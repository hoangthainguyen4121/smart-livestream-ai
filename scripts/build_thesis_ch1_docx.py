"""Build Chapter 1 DOCX and thesis figures from frozen planning docs."""
from __future__ import annotations

import glob
import shutil
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = ROOT / "docs" / "thesis-figures"
SHOT_DIR = ROOT / "docs" / "thesis-screenshots"
OUT_DIR = ROOT / "Document"
TEMPLATE_GLOB = r"D:\Thac Si NTTU\Chuyen de cntt\Document\4.*.docx"

TITLE = (
    "Xây dựng hệ thống hỗ trợ bán hàng thông minh cho livestream "
    "ứng dụng trí tuệ nhân tạo và thị giác máy tính"
)


def _box(ax, xy, w, h, text, fc="#E8F4FD", ec="#1a5276", fontsize=9):
    x, y = xy
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.08",
        linewidth=1.2,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, wrap=True)


def draw_hinh_1_1(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    ax.text(5, 5.5, "Bối cảnh bán hàng qua phiên phát trực tiếp và quá tải bình luận", ha="center", fontsize=12, fontweight="bold")
    _box(ax, (0.4, 3.2), 2.2, 1.4, "Người phát\n trực tiếp\n(trình diễn sản phẩm)", fc="#FDEBD0")
    _box(ax, (3.5, 3.8), 3.0, 0.9, "Nền tảng thương mại\n(Shopee Live, TikTok Shop, …)", fc="#FADBD8")
    _box(ax, (7.2, 3.2), 2.4, 1.4, "Người xem\n(hàng trăm–nghìn\nbình luận)", fc="#D5F5E3")
    for i, msg in enumerate(
        ["Giá?", "Còn hàng?", "Link mua?", "Chốt đơn", "Thử son", "Size?"]
    ):
        ax.text(3.8 + (i % 3) * 1.1, 2.5 - (i // 3) * 0.55, msg, fontsize=7, bbox=dict(boxstyle="round", fc="#FCF3CF", ec="#B7950B"))
    ax.annotate("", xy=(3.5, 4.2), xytext=(2.6, 3.9), arrowprops=dict(arrowstyle="->", color="#566573"))
    ax.annotate("", xy=(6.5, 4.2), xytext=(7.2, 3.9), arrowprops=dict(arrowstyle="->", color="#566573"))
    _box(ax, (2.0, 0.5), 6.0, 1.5, "Áp lực: người phát trực tiếp không kịp phản hồi · thiếu thử sản phẩm trực quan\n· cử chỉ chưa được ghi nhận · các chức năng hỗ trợ triển khai rời rạc", fc="#EBF5FB", fontsize=9)
    ax.text(5, -0.15, "Nguồn: Tổng hợp tác giả, 2026", ha="center", fontsize=8, style="italic")
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def draw_hinh_1_2(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 3.2))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 3)
    ax.axis("off")
    steps = [
        ("1. Bài toán", "Quá tải bình luận\ntrong phiên phát\ntrực tiếp"),
        ("2. Thách thức", "Phản hồi chậm,\nthiếu thử sản phẩm,\ntín hiệu phi ngôn ngữ"),
        ("3. Giải pháp", "Hệ thống hỗ trợ\nbán hàng trên\ntrình duyệt web"),
        ("4. Kiến trúc\n& hiện thực", "Nguyên mẫu tích hợp\ncác chức năng\nnghiệp vụ"),
        ("5. Đánh giá", "Kịch bản thử nghiệm\nvà chỉ số đo lường"),
    ]
    w = 1.85
    for i, (title, body) in enumerate(steps):
        x = 0.3 + i * 2.15
        _box(ax, (x, 0.9), w, 1.5, f"{title}\n\n{body}", fc="#E8DAEF" if i % 2 else "#D6EAF8", fontsize=8)
        if i < len(steps) - 1:
            ax.annotate("", xy=(x + w + 0.08, 1.65), xytext=(x + w + 0.35, 1.65), arrowprops=dict(arrowstyle="->", lw=1.5))
    ax.text(5.5, 2.75, "Quy trình nghiên cứu: bài toán → thách thức → giải pháp → kiến trúc → đánh giá", ha="center", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def draw_hinh_3_1(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.text(5.5, 6.65, "Kiến trúc hiện trạng của nguyên mẫu (triển khai trên trình duyệt web)", ha="center", fontsize=12, fontweight="bold")
    _box(ax, (0.5, 4.6), 10.0, 1.2, "Trình duyệt web — giao diện người dùng (React)\nTrò chuyện · thực tế tăng cường · giỏ hàng · trợ lý bán hàng", fc="#D4EFDF")
    _box(ax, (0.5, 2.8), 10.0, 1.2, "Máy chủ ứng dụng (FastAPI) — trò chuyện · đăng ký khuôn mặt · sự kiện tương tác · chuyển tiếp xử lý ngôn ngữ", fc="#D6EAF8")
    _box(ax, (0.5, 1.5), 4.5, 0.9, "Phân loại ý định bình luận (PhoBERT)", fc="#FCF3CF")
    _box(ax, (5.3, 1.5), 5.2, 0.9, "Nhận diện cử chỉ (MediaPipe) · nhận diện khuôn mặt (InsightFace)", fc="#FADBD8")
    _box(ax, (0.5, 0.2), 10.0, 0.9, "Luồng nghiệp vụ: bình luận → ý định → sản phẩm → phản hồi → giỏ hàng / thanh toán mô phỏng", fc="#EBF5FB", fontsize=8)
    for y1, y2 in [(4.6, 4.0), (2.8, 2.4)]:
        ax.annotate("", xy=(5.5, y2), xytext=(5.5, y1), arrowprops=dict(arrowstyle="->", lw=1.5))
    ax.annotate("", xy=(2.75, 2.4), xytext=(2.75, 2.8), arrowprops=dict(arrowstyle="->", lw=1.5))
    ax.text(5.5, -0.05, "Không bao gồm: ứng dụng desktop · PostgreSQL · cổng thanh toán · OAuth (hướng phát triển, mục 3.10)", ha="center", fontsize=7.5, style="italic")
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def draw_hinh_3_5(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 4.5))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.text(5.5, 4.65, "Quy trình xử lý bình luận của trợ lý bán hàng thông minh", ha="center", fontsize=12, fontweight="bold")
    nodes = [
        (0.3, "Bình luận\nngười xem"),
        (2.0, "Phân loại\ný định\n(PhoBERT + luật)"),
        (4.0, "Xác định\nsản phẩm\nđược nhắc"),
        (6.0, "Sinh câu\nphản hồi"),
        (8.0, "Gợi ý thao tác\nmua hàng"),
        (9.8, "Giỏ hàng /\nthanh toán\nmô phỏng"),
    ]
    w = 1.55
    for i, (x, label) in enumerate(nodes):
        _box(ax, (x, 1.8), w, 1.3, label, fc="#D5F5E3" if i % 2 else "#D6EAF8", fontsize=8)
        if i < len(nodes) - 1:
            ax.annotate("", xy=(x + w + 0.05, 2.45), xytext=(x + w + 0.35, 2.45), arrowprops=dict(arrowstyle="->", lw=1.5))
    _box(ax, (2.0, 0.35), 6.5, 0.85, "Ghi nhận cử chỉ và thống kê tương tác trên bảng sự kiện AI (lưu tạm trong phiên)", fc="#FEF9E7", fontsize=8)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def set_run_font(run, name="Times New Roman", size=13, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc, text, style="Normal", bold=False, size=13, align=None, space_after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    """level 1 = chapter, 2 = section 1.1, 3 = subsection 1.1.1"""
    sizes = {1: 14, 2: 13, 3: 13}
    bold = True
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 13), bold=bold)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_figure(doc, image_path: Path, caption: str, width_cm=14):
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=12, bold=True)
    cap.paragraph_format.space_after = Pt(12)


def add_table_1_1(doc):
    add_paragraph(doc, "Bảng 1.1. Mục tiêu nghiên cứu và tiêu chí đánh giá thành công", bold=True, size=13)
    rows = [
        ("STT", "Mục tiêu / năng lực", "Tiêu chí đánh giá (định tính)"),
        (
            "1",
            "Thử sản phẩm bằng thực tế tăng cường trên trình duyệt web",
            "Người xem quan sát hiệu ứng thử sản phẩm (kính, trang điểm, bộ lọc) trên luồng hình ảnh; sản phẩm được ghim thì hiệu ứng thay đổi tương ứng",
        ),
        (
            "2",
            "Đăng ký và lưu trữ đặc trưng khuôn mặt qua trình duyệt",
            "Người dùng thực hiện đăng ký đa tư thế trên giao diện web; vectơ đặc trưng được lưu cục bộ phục vụ nhận diện",
        ),
        (
            "3",
            "Ghi nhận cử chỉ và hiển thị trên bảng sự kiện AI",
            "Cử chỉ giơ tay hoặc giơ ngón cái được ghi nhận và hiển thị; tần suất lặp được giới hạn",
        ),
        (
            "4",
            "Phân loại ý định trong bình luận của người xem",
            "Bình luận mang ý định mua bán được gán nhãn; hệ thống cho biết kết quả từ mô hình PhoBERT hoặc phương án dự phòng dựa trên luật",
        ),
        (
            "5",
            "Xác định sản phẩm được nhắc đến trong bình luận",
            "Liên kết nội dung bình luận với sản phẩm trong danh mục mười mục của nguyên mẫu",
        ),
        (
            "6",
            "Hỗ trợ phản hồi và gợi ý thao tác mua hàng",
            "Trợ lý bán hàng thông minh sinh câu trả lời và gợi ý thao tác (thêm giỏ hàng, mở quy trình thanh toán) theo ý định nhận diện",
        ),
        (
            "7",
            "Trò chuyện trực tuyến đồng bộ giữa nhiều thiết bị xem",
            "Nội dung trò chuyện được đồng bộ giữa nhiều thẻ trình duyệt trong cùng phiên nguyên mẫu (WebSocket)",
        ),
        (
            "8",
            "Mô phỏng giỏ hàng và quy trình thanh toán",
            "Thực hiện thêm sản phẩm, mở biểu mẫu thanh toán mô phỏng và nhận mã đơn hàng trong phiên trình duyệt",
        ),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=12, bold=(i == 0))
    add_paragraph(doc, "Nguồn: Tổng hợp tác giả theo phạm vi hiện thực nguyên mẫu thử nghiệm.", size=12, space_after=12)


def add_table_1_2(doc):
    add_paragraph(doc, "Bảng 1.2. Phạm vi nghiên cứu: trong phạm vi và ngoài phạm vi", bold=True, size=13)
    rows = [
        ("Phân loại", "Nội dung"),
        (
            "Trong phạm vi",
            "Nguyên mẫu hỗ trợ bán hàng trong phiên phát trực tiếp, vận hành trên trình duyệt web, gồm thử sản phẩm trực quan, trò chuyện, trợ lý bán hàng thông minh và mô phỏng quy trình mua hàng",
        ),
        (
            "",
            "Đăng ký khuôn mặt qua trình duyệt; danh mục mười sản phẩm mô phỏng; huấn luyện và sử dụng mô hình PhoBERT phân loại ý định trên tập dữ liệu chủ yếu tổng hợp (tám bình luận thực tế)",
        ),
        (
            "",
            "Triển khai thử nghiệm cục bộ hoặc bằng Docker; dữ liệu trò chuyện và sự kiện lưu tạm trong bộ nhớ, mất khi khởi động lại dịch vụ",
        ),
        (
            "Ngoài phạm vi",
            "Cơ sở dữ liệu quan hệ PostgreSQL, xác thực OAuth/JWT, chatbot hỏi đáp cố định, cổng thanh toán thật, triển khai Kubernetes",
        ),
        (
            "",
            "Ứng dụng desktop gốc, cửa sổ OpenCV — đã loại khỏi phạm đề tài có chủ đích",
        ),
        (
            "",
            "Lưu trữ lâu dài thông tin cá nhân và đơn hàng trong cơ sở dữ liệu",
        ),
        (
            "",
            "Chỉ số khung hình/giây của thực tế tăng cường, kết quả kiểm thử tự động chi tiết — trình bày tại Chương 4, không nêu tại Chương 1",
        ),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=12, bold=(i == 0))
    add_paragraph(doc, "Nguồn: Tổng hợp tác giả theo phạm vi triển khai thực tế của nguyên mẫu.", size=12, space_after=12)


def build_chapter_1_body(doc, fig_dir: Path):
    add_heading(doc, "CHƯƠNG 1", level=1)
    add_heading(doc, "GIỚI THIỆU", level=1)

    add_heading(doc, "1.1. Đặt vấn đề", level=2)
    add_heading(doc, "1.1.1. Bối cảnh thực tiễn bán hàng qua phiên phát trực tiếp", level=3)
    add_paragraph(
        doc,
        "Trong bối cảnh thương mại điện tử tại Việt Nam, bán hàng qua phiên phát trực tiếp trên "
        "các nền tảng như Shopee Live, TikTok Shop và Facebook Live đã trở thành hình thức phân phối "
        "phổ biến. Người phát trực tiếp trình bày sản phẩm trước camera; người xem theo dõi và tương "
        "tác bằng bình luận về giá cả, tình trạng hàng hóa, liên kết mua và ý định đặt hàng. "
        "Khối lượng bình luận trong một phiên có thể lên đến hàng trăm, thậm chí hàng nghìn lượt.",
    )
    add_paragraph(
        doc,
        "Từ góc độ nghiệp vụ, bài toán đặt ra là làm sao hỗ trợ người phát trực tiếp xử lý tương tác "
        "mua bán kịp thời, đồng thời giúp người xem hình dung sản phẩm và thực hiện thao tác mua "
        "hàng mà không phải chờ phản hồi thủ công cho từng bình luận.",
    )
    add_figure(
        doc,
        fig_dir / "hinh_1_1_livestream_context.png",
        "Hình 1.1. Bối cảnh bán hàng qua phiên phát trực tiếp và quá tải bình luận",
    )

    add_heading(doc, "1.1.2. Thách thức trong xử lý bình luận và chốt đơn", level=3)
    add_paragraph(
        doc,
        "Thách thức thứ nhất là quá tải bình luận. Người phát trực tiếp khó đọc và trả lời đầy đủ "
        "các bình luận mang ý định hỏi giá, xin thông tin mua hàng hoặc chốt đơn. Khi bình luận "
        "bị bỏ qua, người xem có thể rời phiên hoặc trì hoãn quyết định mua.",
    )
    add_paragraph(
        doc,
        "Thách thức thứ hai liên quan đến trải nghiệm trực quan. Người xem thường cần hình dung "
        "sản phẩm (ví dụ kính mắt, son môi, phụ kiện) ngay trên luồng hình ảnh, thay vì chỉ nghe "
        "mô tả bằng lời.",
    )
    add_paragraph(
        doc,
        "Thách thức thứ ba nằm ở tín hiệu phi ngôn ngữ. Người xem có thể giơ tay hoặc giơ ngón "
        "cái để thể hiện quan tâm, song các tín hiệu này hiếm khi được hệ thống ghi nhận nếu chỉ "
        "tập trung xử lý văn bản bình luận.",
    )

    add_heading(doc, "1.1.3. Hạn chế của các giải pháp triển khai rời rạc", level=3)
    add_paragraph(
        doc,
        "Hiện nay, chức năng thử sản phẩm bằng thực tế tăng cường, trả lời bình luận tự động và "
        "quản lý giỏ hàng thường được triển khai tách biệt. Thực tế tăng cường có thể không liên "
        "kết với nội dung bình luận; câu trả lời tự động có thể không dẫn tới thao tác mua hàng; "
        "cử chỉ người xem có thể không xuất hiện trên cùng một giao diện theo dõi. Do đó, khó "
        "đánh giá hiệu quả hỗ trợ bán hàng theo một luồng nghiệp vụ thống nhất trong phiên phát trực tiếp.",
    )
    add_paragraph(
        doc,
        "Xuất phát từ các hạn chế trên, luận văn hướng tới việc xây dựng nguyên mẫu thử nghiệm "
        "tích hợp các chức năng hỗ trợ trên một giao diện trình duyệt web, thay vì phụ thuộc "
        "ứng dụng desktop riêng.",
    )

    add_heading(doc, "1.2. Mục tiêu nghiên cứu và phát triển hệ thống", level=2)
    add_heading(doc, "1.2.1. Mục tiêu tổng quát", level=3)
    add_paragraph(
        doc,
        f"Mục tiêu tổng quát của luận văn là {TITLE.lower()}. "
        "Cụ thể, nghiên cứu xây dựng nguyên mẫu hỗ trợ người phát trực tiếp và người xem trong "
        "phiên bán hàng, kết hợp thị giác máy tính và xử lý ngôn ngữ tự nhiên trên nền tảng web.",
    )
    add_paragraph(
        doc,
        "Phạm vi đóng góp tập trung vào hệ thống tích hợp phục vụ bài toán nghiệp vụ nêu trên, "
        "không nhằm nghiên cứu độc lập toàn bộ quy trình huấn luyện mô hình ngôn ngữ. Mô hình "
        "PhoBERT được sử dụng như một công cụ phân loại ý định trong bình luận của người xem.",
    )

    add_heading(doc, "1.2.2. Mục tiêu cụ thể", level=3)
    add_paragraph(
        doc,
        "Căn cứ bài toán và thách thức đã nêu, mục tiêu cụ thể được xác định theo các chức năng "
        "hệ thống cần đáp ứng: (1) thử sản phẩm bằng thực tế tăng cường trên trình duyệt web; "
        "(2) đăng ký và sử dụng đặc trưng khuôn mặt qua trình duyệt; (3) ghi nhận cử chỉ và hiển thị "
        "trên bảng sự kiện AI; (4) phân loại ý định trong bình luận; (5) xác định sản phẩm được "
        "nhắc đến; (6) hỗ trợ phản hồi và gợi ý thao tác mua hàng thông qua trợ lý bán hàng thông minh; "
        "(7) trò chuyện trực tuyến đồng bộ giữa nhiều thiết bị xem; (8) mô phỏng giỏ hàng và quy trình thanh toán.",
    )

    add_heading(doc, "1.2.3. Tiêu chí đánh giá thành công", level=3)
    add_paragraph(
        doc,
        "Tại Chương 1, tiêu chí đánh giá được nêu ở mức định tính, gắn với từng chức năng trong "
        "Bảng 1.1. Các chỉ số đo lường (độ chính xác phân loại ý định, khung hình/giây của thực tế "
        "tăng cường, kết quả kiểm thử tự động) chỉ được trình bày tại Chương 4 sau khi tiến hành "
        "thí nghiệm; luận văn không nêu số liệu chưa được đo đạc tại chương giới thiệu.",
    )
    add_table_1_1(doc)

    add_heading(doc, "1.3. Phạm vi nghiên cứu", level=2)
    add_heading(doc, "1.3.1. Đối tượng nghiên cứu", level=3)
    add_paragraph(
        doc,
        "Đối tượng nghiên cứu là nguyên mẫu thử nghiệm (PoC) hệ thống hỗ trợ bán hàng trong phiên "
        "phát trực tiếp. Nguyên mẫu tập trung vào các hoạt động: hỗ trợ người xem thử sản phẩm "
        "trực quan; ghi nhận cử chỉ và đặc trưng khuôn mặt; đọc hiểu bình luận để phân loại ý định "
        "và xác định sản phẩm; sinh phản hồi cùng gợi ý thao tác mua hàng; mô phỏng giỏ hàng và "
        "quy trình thanh toán. Toàn bộ luồng chính được trình diễn trên trình duyệt web.",
    )
    add_paragraph(
        doc,
        "Nguyên mẫu không bao gồm ứng dụng desktop gốc hay cửa sổ hiển thị video độc lập ngoài "
        "trình duyệt.",
    )

    add_heading(doc, "1.3.2. Dữ liệu sử dụng", level=3)
    add_paragraph(
        doc,
        "Về dữ liệu sản phẩm, nguyên mẫu sử dụng danh mục mười sản phẩm mô phỏng phục vụ thử nghiệm "
        "chức năng bán hàng. Về dữ liệu phân loại ý định, nguyên mẫu sử dụng tập huấn luyện cho mô "
        "hình PhoBERT, trong đó phần lớn mẫu được tổng hợp và chỉ có tám bình luận tiếng Việt "
        "thu thập từ phiên phát trực tiếp thực tế. Hạn chế về quy mô dữ liệu thực được ghi nhận "
        "ngay trong phạm vi nghiên cứu.",
    )
    add_paragraph(
        doc,
        "Dữ liệu nhập trong quy trình thanh toán mô phỏng (họ tên, số điện thoại, địa chỉ) chỉ tồn "
        "tại trong phiên trình duyệt hiện tại, không lưu trữ lâu dài.",
    )

    add_heading(doc, "1.3.3. Phạm vi triển khai", level=3)
    add_paragraph(
        doc,
        "Nguyên mẫu phục vụ mục đích nghiên cứu và trình diễn, triển khai thử nghiệm trên môi "
        "trường cục bộ hoặc bằng Docker. Dữ liệu trò chuyện và sự kiện tương tác được lưu tạm "
        "trong bộ nhớ và mất khi khởi động lại dịch vụ. Luận văn không triển khai cơ sở dữ liệu "
        "đơn hàng bền vững, cơ chế xác thực người dùng OAuth/JWT hay kết nối cổng thanh toán thực.",
    )
    add_table_1_2(doc)

    add_heading(doc, "1.4. Phương pháp thực hiện", level=2)
    add_heading(doc, "1.4.1. Phương pháp nghiên cứu ứng dụng", level=3)
    add_paragraph(
        doc,
        "Luận văn áp dụng phương pháp nghiên cứu ứng dụng theo trình tự: khảo sát bài toán nghiệp "
        "vụ trong phiên phát trực tiếp; đề xuất giải pháp tích hợp; thiết kế kiến trúc; hiện thực "
        "nguyên mẫu; đánh giá bằng kịch bản thử nghiệm. Các nhận định về chức năng hệ thống được "
        "đối chiếu với mã nguồn và kết quả kiểm thử trước khi trình bày tại các chương sau.",
    )

    add_heading(doc, "1.4.2. Quy trình nghiên cứu từ bài toán đến đánh giá", level=3)
    add_paragraph(
        doc,
        "Quy trình nghiên cứu bám sát chuỗi logic: (1) xác định bài toán quá tải bình luận và "
        "nhu cầu thử sản phẩm; (2) phân tích thách thức về phản hồi, trải nghiệm trực quan và tín "
        "hiệu phi ngôn ngữ; (3) đề xuất nguyên mẫu tích hợp trên trình duyệt web (Chương 3); "
        "(4) hiện thực và mô tả các chức năng (Chương 4); (5) đánh giá bằng kịch bản thử nghiệm "
        "và chỉ số đo lường có căn cứ. Chương 1 dừng ở bước xác định bài toán, mục tiêu và phương pháp.",
    )
    add_figure(
        doc,
        fig_dir / "hinh_1_2_research_process.png",
        "Hình 1.2. Quy trình nghiên cứu: bài toán → thách thức → giải pháp → kiến trúc → đánh giá",
    )

    add_heading(doc, "1.4.3. Lựa chọn công nghệ (mức tổng quan)", level=3)
    add_paragraph(
        doc,
        "Công nghệ chỉ được lựa chọn sau khi xác định yêu cầu nghiệp vụ. Vì nguyên mẫu cần tương "
        "tác thời gian thực trên trình duyệt, hệ thống sử dụng React/TypeScript cho giao diện và "
        "FastAPI/Python cho máy chủ ứng dụng. Vì bình luận mang ngữ cảnh mua bán tiếng Việt, hệ "
        "thống sử dụng PhoBERT cho phân loại ý định, kết hợp phương án dự phòng dựa trên luật. "
        "Vì cần thử sản phẩm và ghi nhận cử chỉ trên luồng hình ảnh, nguyên mẫu sử dụng thực tế "
        "tăng cường phía trình duyệt cùng MediaPipe và InsightFace cho nhận diện cử chỉ và khuôn mặt. "
        "Docker hỗ trợ triển khai thử nghiệm cục bộ. Chi tiết phiên bản phần mềm và cấu hình phần cứng "
        "được trình bày tại mục 4.1.",
    )

    add_heading(doc, "1.5. Cấu trúc tiểu luận", level=2)
    add_paragraph(
        doc,
        "Luận văn gồm năm chương, bám theo trục: bài toán thực tế → thách thức → giải pháp đề xuất → "
        "thiết kế hệ thống → hiện thực hệ thống → đánh giá. Trung tâm là hệ thống hỗ trợ bán hàng "
        "thông minh trong phiên phát trực tiếp, tích hợp hai trụ cột thị giác máy tính và trí tuệ "
        "nhân tạo (xử lý ngôn ngữ tự nhiên) theo đúng tên đề tài.",
    )
    add_paragraph(
        doc,
        "Chương 1 — Giới thiệu: bối cảnh, tính cấp thiết, mục tiêu, phạm vi, phương pháp nghiên cứu "
        "và cấu trúc luận văn.",
    )
    add_paragraph(
        doc,
        "Chương 2 — Cơ sở lý thuyết và nghiên cứu liên quan: tổng quan thương mại phiên phát trực tiếp "
        "(2.1); những thách thức trong hỗ trợ bán hàng trực tuyến (2.2); thị giác máy tính trong hỗ trợ "
        "bán hàng — nhận diện khuôn mặt, nhận dạng cử chỉ, thực tế tăng cường (2.3); xử lý ngôn ngữ "
        "tự nhiên trong bình luận người xem — phân loại ý định, PhoBERT, TF-IDF (2.4); trợ lý bán hàng "
        "thông minh (2.5); các nghiên cứu liên quan (2.6); công nghệ sử dụng trong hệ thống (2.7).",
    )
    add_paragraph(
        doc,
        "Chương 3 — Phân tích và thiết kế hệ thống: phân tích yêu cầu (3.1); kiến trúc tổng thể (3.2); "
        "thiết kế xử lý hình ảnh (3.3); thiết kế xử lý bình luận (3.4); thiết kế trợ lý bán hàng "
        "thông minh (3.5); thiết kế quy trình mua hàng mô phỏng (3.6); kiến trúc triển khai (3.7); "
        "tóm tắt chương (3.8). Chương này không nêu hướng phát triển.",
    )
    add_paragraph(
        doc,
        "Chương 4 — Xây dựng và đánh giá hệ thống: môi trường thực nghiệm (4.1); giao diện hệ thống "
        "(4.2); đánh giá nhận diện khuôn mặt (4.3), nhận dạng cử chỉ (4.4), phân tích ý định bình luận "
        "(4.5), trợ lý bán hàng thông minh (4.6), quy trình mua hàng mô phỏng (4.7); đánh giá tổng thể "
        "(4.8).",
    )
    add_paragraph(
        doc,
        "Chương 5 — Kết luận và hướng phát triển: kết quả, đóng góp, hạn chế; hướng phát triển chỉ "
        "trình bày tại mục 5.4.",
    )
    add_paragraph(
        doc,
        "Phụ lục: mã nguồn, hướng dẫn vận hành nguyên mẫu, ảnh giao diện bổ sung, cấu hình huấn luyện "
        "PhoBERT và nhật ký kiểm thử.",
        space_after=12,
    )


def fill_cover(doc):
    filled = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t in ("Tên đề tài:", "Tên đề tài") and filled < 2:
            if i + 1 < len(doc.paragraphs):
                np = doc.paragraphs[i + 1]
                if "(Bold, size" not in np.text:
                    np.clear()
                    r = np.add_run(TITLE)
                    set_run_font(r, size=16, bold=True)
                    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    filled += 1


def remove_after_chapter_sample(doc):
    """Remove template sample from CHƯƠNG 1 placeholder onward."""
    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "CHƯƠNG 1 (bold, size 16-)":
            start_idx = i
            break
    if start_idx is not None:
        paras = list(doc.paragraphs)
        for p in paras[start_idx:]:
            p._element.getparent().remove(p._element)


def main():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    SHOT_DIR.mkdir(parents=True, exist_ok=True)

    print("Generating figures...")
    draw_hinh_1_1(FIG_DIR / "hinh_1_1_livestream_context.png")
    draw_hinh_1_2(FIG_DIR / "hinh_1_2_research_process.png")
    draw_hinh_3_1(FIG_DIR / "hinh_3_1_as_is_architecture.png")
    draw_hinh_3_5(FIG_DIR / "hinh_3_5_sales_assistant_pipeline.png")

    templates = glob.glob(TEMPLATE_GLOB)
    if not templates:
        raise FileNotFoundError("University template DOCX not found")
    out_path = OUT_DIR / "Tieu luan - Smart Livestream - Chuong 1.docx"
    shutil.copy2(templates[0], out_path)

    doc = Document(str(out_path))
    fill_cover(doc)
    remove_after_chapter_sample(doc)
    build_chapter_1_body(doc, FIG_DIR)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")
    print(f"Figures in {FIG_DIR}")


if __name__ == "__main__":
    main()
