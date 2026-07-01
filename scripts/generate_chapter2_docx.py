# -*- coding: utf-8 -*-
"""Generate Chapter 2 DOCX — Batch Writing Mode.

Review unit: batch clusters (e.g. §2.2–2.3). Output: Document/ only until batch locked.
Manuscript sync: --sync-manuscript when batch is locked.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = ROOT / "docs" / "thesis-figures"
OUT = ROOT / "Document" / "Tieu luan - Smart Livestream - Chuong 2.docx"
OUT_REVIEW = ROOT / "thesis_review_package" / "manuscript" / OUT.name

# Chapter lock status
CHAPTER_LOCKED = True
REVIEW_NOTE = "docs/chapter2/REVIEW_NOTE_BATCH3_2_6_2_7.md"


def set_run_font(run, name="Times New Roman", size=13, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_para(doc, text, size=13, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, size=13, bold=True):
    return add_para(doc, text, size=size, bold=bold, space_after=8)


def add_body(doc, text):
    return add_para(doc, text, size=13, bold=False, space_after=6)


def add_caption(doc, text):
    p = add_para(doc, text, size=13, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    for run in p.runs:
        run.italic = True
    return p


def add_figure(doc, filename, caption, width_cm=14):
    path = FIG_DIR / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=12)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=12)
    doc.add_paragraph()
    return table


def add_chapter_intro(doc):
    add_heading(doc, "CHƯƠNG 2", size=14)
    add_heading(doc, "CƠ SỞ LÝ THUYẾT VÀ NGHIÊN CỨU LIÊN QUAN", size=14)
    add_body(
        doc,
        "Chương này trình bày cơ sở lý thuyết phục vụ thiết kế hệ thống hỗ trợ bán hàng trong phiên "
        "phát trực tiếp. Nội dung được tổ chức theo trục nghiệp vụ: bối cảnh thương mại trực tuyến, "
        "các thách thức vận hành, hai trụ cột công nghệ (thị giác máy tính và xử lý ngôn ngữ tự nhiên), "
        "khái niệm trợ lý bán hàng, tổng quan nghiên cứu liên quan và tóm tắt công nghệ triển khai. "
        "Chương tập trung vào khung lý thuyết; mô tả thiết kế chi tiết và kết quả thử nghiệm được trình bày "
        "lần lượt tại Chương 3 và Chương 4.",
    )


def add_section_2_1(doc):
    """§2.1 — full section (2.1.1 + 2.1.2). Review at level 2.1."""
    add_heading(doc, "2.1. Tổng quan về thương mại qua phiên phát trực tiếp")

    # --- 2.1.1 (locked content rev.1) ---
    add_heading(doc, "2.1.1. Bối cảnh học thuật và mô hình tương tác trong phiên bán hàng trực tiếp")
    add_body(
        doc,
        "Thương mại qua phiên phát trực tiếp (livestream commerce) là mô hình bán hàng trong đó trình bày sản phẩm, "
        "tương tác với khách hàng và thúc đẩy quyết định mua diễn ra đồng thời trên cùng một phiên thời gian thực. "
        "Khác với thương mại điện tử bất đồng bộ, phiên phát trực tiếp coi tương tác trực tiếp là thành phần cấu thành "
        "của giá trị mua bán. Tiểu mục này làm rõ cơ chế tương tác ở mức lý thuyết, làm cơ sở cho thiết kế hệ thống "
        "ở các chương sau.",
    )
    add_body(
        doc,
        "Wongkitrungrueng và Assarut (2020) — đã được trích dẫn tại Chương 1 — cho thấy livestream trên nền "
        "tảng thương mại xã hội không chỉ truyền tải thông tin sản phẩm mà còn tạo giá trị cảm nhận ở ba chiều: "
        "công cụ (thông tin hữu ích), cảm xúc (trải nghiệm thú vị) và biểu tượng (danh tiếng người bán). Ba chiều "
        "này tác động gián tiếp lên ý định mua thông qua tin cậy sản phẩm và tin cậy người bán. Khi tin cậy được xây dựng "
        "qua tương tác thời gian thực, việc phản hồi chậm hoặc thiếu phản hồi trong phiên có thể làm suy giảm cả hai chiều "
        "tin cậy — tức làm giảm hiệu quả của chính cơ chế mà mô hình livestream commerce dựa vào. Chương 1 đã phản ánh "
        "áp lực này dưới góc vận hành; tại đây, áp lực được neo vào logic lý thuyết về engagement, làm cơ sở cho việc "
        "xem bình luận như tín hiệu nghiệp vụ cần được hỗ trợ, không phải luồng phụ trợ tùy chọn.",
    )
    add_body(
        doc,
        "Lu và Chen (2021) bổ sung khía cạnh cấu trúc tương tác của dịch vụ phát trực tiếp xã hội: người xem vừa "
        "tiêu thụ nội dung video, vừa đồng sản xuất engagement qua bình luận, biểu tượng cảm xúc và câu hỏi liên quan "
        "sản phẩm. Phiên bán hàng vì vậy có ít nhất hai luồng song song — luồng trình diễn do người phát kiểm soát và "
        "luồng phản hồi do khán giả khởi tạo — tạo thành một quy trình hỗ trợ bán hàng đặc thù mà mô hình thương mại "
        "điện tử truyền thống (trang sản phẩm, giỏ hàng tĩnh) không mô phỏng đầy đủ. Khác với Chương 1 nhấn mạnh quy mô "
        "bình luận, tiểu mục này nhấn mạnh tính đồng thời và công khai của tương tác: mỗi câu trả lời (hoặc không trả lời) "
        "đều diễn ra trước cộng đồng người xem, khiến chất lượng phản hồi gắn trực tiếp với trải nghiệm mua hàng của "
        "nhiều người cùng lúc.",
    )
    add_body(
        doc,
        "Từ các nghiên cứu trên, có thể nhận thấy rằng livestream commerce lấy tương tác làm trung tâm: tin cậy và "
        "quyết định mua phụ thuộc chất lượng và kịp thời của giao tiếp trong phiên. Trong phạm vi luận văn này, tiền đề "
        "đó giải thích vì sao Chương 2 xuất phát từ nghiệp vụ thay vì công nghệ, đồng thời khẳng định bình luận là tín "
        "hiệu gắn với cơ chế tin cậy, không phải dữ liệu phụ trợ tách rời quy trình bán hàng.",
    )
    add_body(
        doc,
        "Kết luận lại, khi tương tác thời gian thực vừa tạo tin cậy vừa gây áp lực vận hành, nhu cầu xây dựng hệ thống "
        "hỗ trợ bán hàng thông minh trở nên hợp lý: hệ thống cần ghi nhận, phân loại và hỗ trợ xử lý luồng tương tác mà "
        "các nghiên cứu trên chưa giải quyết ở mức công cụ cho người phát. Tiểu mục 2.1.2 sẽ chuyển sang mức vận hành — "
        "vai trò người phát, phân loại bình luận và áp lực xử lý đồng thời — làm cầu nối tới mục 2.2 và 3.1.",
    )

    # --- 2.1.2 ---
    add_heading(doc, "2.1.2. Vai trò vận hành, loại bình luận và quy trình hỗ trợ bán hàng trong phiên")
    add_body(
        doc,
        "Tiểu mục 2.1.1 đã chỉ ra rằng tương tác thời gian thực là trục của livestream commerce. Để chuyển tiền đề "
        "lý thuyết sang đối tượng thiết kế hệ thống, cần mô tả phiên bán hàng ở mức vận hành: ai tham gia, bình luận "
        "mang ý nghĩa gì và áp lực xử lý diễn ra như thế nào khi hai luồng trình diễn và phản hồi chạy song song.",
    )
    add_body(
        doc,
        "Trong phiên bán hàng trực tiếp, người phát trực tiếp không chỉ giới thiệu sản phẩm mà còn là điểm tiếp nhận "
        "phản hồi, duy trì nhịp phiên và quyết định cuối về chốt đơn hay xử lý khiếu nại. Wongkitrungrueng và Assarut "
        "(2020) cho thấy người bán đóng vai trò then chốt trong hình thành tin cậy; Lu và Chen (2021) nhấn mạnh người "
        "xem đồng thời là người tạo engagement. Kết hợp hai góc nhìn, có thể xem người phát như trung tâm điều phối: "
        "vừa trình bày sản phẩm, vừa phải đọc và phản hồi luồng bình luận, trong khi người xem vừa theo dõi vừa gửi "
        "tín hiệu mua bán qua nhiều kênh — bình luận văn bản, cử chỉ trên camera và nhu cầu thử sản phẩm trực quan.",
    )
    add_body(
        doc,
        "Bình luận trong phiên không đồng nhất về mục đích. Ở mức nghiệp vụ — trước khi phân loại chi tiết tại mục 2.4 — "
        "có thể nhóm thành các nhóm chính: hỏi thông tin sản phẩm (giá, biến thể, tồn kho, vận chuyển, khuyến mãi); "
        "yêu cầu hành động mua (xin liên kết, ý định chốt đơn); trao đổi xã giao hoặc nội dung không liên quan trực tiếp "
        "đến mua bán. Phân biệt này quan trọng vì không phải mọi bình luận đều đòi hỏi phản hồi cùng mức ưu tiên: câu hỏi "
        "về giá hoặc chốt đơn thường gắn trực tiếp với chuyển đổi bán hàng, trong khi bình luận xã giao có thể được hoãn "
        "xử lý hoặc bỏ qua.",
    )
    add_body(
        doc,
        "Áp lực vận hành nằm ở chỗ người phát phải xử lý đồng thời nhiều bình luận trong khi vẫn duy trì luồng trình "
        "diễn. Chương 1 đã mô tả khối lượng bình luận có thể lên hàng trăm đến hàng nghìn lượt; kết hợp với luận điểm "
        "§2.1.1, áp lực đó không chỉ là vấn đề khối lượng mà còn là rủi ro suy giảm tin cậy khi câu hỏi mang ý định mua "
        "bị bỏ sót. Ngoài bình luận văn bản, người xem có thể thể hiện quan tâm bằng cử chỉ hoặc thử sản phẩm trên "
        "camera — các tín hiệu này cũng thuộc quy trình hỗ trợ bán hàng nhưng dễ bị bỏ qua nếu người phát chỉ tập trung "
        "theo dõi bình luận.",
    )
    add_body(
        doc,
        "Hình 2.1 mô tả quy trình khái niệm gắn các thành phần trên: người xem tương tác qua bình luận, thử sản phẩm "
        "và tín hiệu hình ảnh; hệ thống hỗ trợ ghi nhận và phân tích; trợ lý bán hàng đề xuất phản hồi hoặc chuyển tiếp "
        "cho người phát; quy trình mua hàng được kích hoạt khi có ý định rõ ràng. Hình thể hiện luồng nghiệp vụ tổng thể "
        "mà nguyên mẫu hướng tới, chưa đi vào chi tiết thuật toán hay triển khai phần mềm.",
    )
    add_figure(
        doc,
        "hinh_2_1_quy_trinh_ho_tro_ban_hang.png",
        "Hình 2.1. Quy trình hỗ trợ bán hàng trong phiên phát trực tiếp",
    )
    add_body(
        doc,
        "Dựa trên các phân tích trên, có thể nhận thấy rằng quy trình hỗ trợ bán hàng trong phiên đòi hỏi một lớp trung "
        "gian ghi nhận đa kênh tín hiệu, phân biệt mức ưu tiên bình luận và hỗ trợ người phát ra quyết định — thay vì để "
        "người phát xử lý thủ công toàn bộ luồng tương tác. Trong phạm vi luận văn này, mô hình vận hành đó là khung tham "
        "chiếu cho nguyên mẫu: mọi chức năng hệ thống đều phục vụ cùng một quy trình phiên, không triển khai rời rạc theo "
        "từng công cụ riêng lẻ.",
    )
    add_body(
        doc,
        "Kết luận, ý nghĩa của tiểu mục này đối với hệ thống được đề xuất là xác định đối tượng cần được hỗ trợ (người "
        "phát trong phiên), dữ liệu đầu vào (bình luận và tín hiệu phi ngôn ngữ) cùng quy trình nghiệp vụ tổng thể thể hiện "
        "ở Hình 2.1 — trước khi formal hóa thách thức cụ thể. Mục 2.2 sẽ chuyển mô hình vận hành này thành bốn nhóm thách "
        "thức, làm cơ sở trực tiếp cho phân tích yêu cầu tại mục 3.1.",
    )


def add_section_2_2(doc):
    add_heading(doc, "2.2. Những thách thức trong hỗ trợ bán hàng trực tuyến")
    add_body(
        doc,
        "Mục 2.1 đã mô tả mô hình vận hành và quy trình hỗ trợ bán hàng trong phiên phát trực tiếp. Mục 2.2 "
        "formal hóa các thách thức nghiệp vụ phát sinh từ mô hình đó — độc lập với lựa chọn công nghệ cụ thể — "
        "nhằm làm cơ sở xác định yêu cầu hệ thống tại mục 3.1. Bốn nhóm thách thức (T1–T4) được xây dựng "
        "xuyên suốt Chương 1 và mục 2.1; tại đây, mỗi nhóm được phân tích theo hệ quả đối với người phát và "
        "người xem trong cùng một phiên.",
    )
    add_body(
        doc,
        "Thách thức T1 — quá tải bình luận: Người phát trực tiếp phải vừa trình bày sản phẩm vừa đọc và phản hồi "
        "hàng loạt bình luận xuất hiện theo thời gian thực. Khi tốc độ bình luận vượt khả năng xử lý thủ công, "
        "các câu hỏi về giá, tồn kho hoặc chốt đơn có thể bị bỏ sót, làm giảm trải nghiệm và cơ hội chuyển đổi. "
        "Wongkitrungrueng và Assarut (2020) nhấn mạnh vai trò của tương tác thời gian thực trong hình thành tin cậy; "
        "Yin và cộng sự (2023) chỉ ra gánh nặng điều phối bình luận trong bán hàng trực tuyến là yếu tố vận hành "
        "đáng kể. T1 là thách thức trực tiếp nhất từ mô hình hai luồng song song đã nêu tại mục 2.1.",
    )
    add_body(
        doc,
        "Thách thức T2 — thiếu thử sản phẩm trực quan: Người xem khó hình dung sản phẩm chỉ qua mô tả bằng lời, "
        "đặc biệt với mỹ phẩm, phụ kiện thời trang hoặc kính mắt. Hoffmann và Mai (2022), trong tổng quan có hệ thống "
        "về mua sắm bằng thực tế tăng cường, cho thấy trải nghiệm thử sản phẩm ảo giúp giảm rủi ro cảm nhận và hỗ trợ "
        "quyết định mua. Tuy nhiên, phần lớn nghiên cứu được tổng hợp triển khai AR trên kênh tách rời phiên phát "
        "trực tiếp; việc gắn thử sản phẩm với sản phẩm đang được giới thiệu trong cùng phiên vẫn còn khoảng trống "
        "thực tiễn.",
    )
    add_body(
        doc,
        "Thách thức T3 — tín hiệu phi ngôn ngữ: Ngoài bình luận văn bản, người xem có thể thể hiện quan tâm bằng "
        "cử chỉ (ví dụ giơ tay, giơ ngón cái) hoặc xuất hiện trên khung hình. Các tín hiệu này mang thông tin hỗ trợ "
        "bán hàng nhưng thường không được đưa vào quy trình xử lý nếu hệ thống chỉ tập trung vào văn bản. Mục 2.1 đã "
        "nêu rằng quy trình hỗ trợ cần ghi nhận đa kênh; T3 cụ thể hóa khoảng trống đó ở chiều hình ảnh và hành vi "
        "phi ngôn ngữ.",
    )
    add_body(
        doc,
        "Thách thức T4 — giải pháp triển khai rời rạc: Thử sản phẩm bằng thực tế tăng cường, chatbot trả lời bình "
        "luận và quản lý giỏ hàng thường được phát triển tách biệt, khó đánh giá như một hệ thống thống nhất phục vụ "
        "cùng một phiên bán hàng. Chương 1 đã phân tích hạn chế này ở mức thực tiễn; tại mục 2.2, T4 được đặt như "
        "thách thức kiến trúc nghiệp vụ: các chức năng hỗ trợ cần cùng phục vụ quy trình thể hiện ở Hình 2.1, thay vì "
        "tồn tại như công cụ độc lập.",
    )
    add_body(
        doc,
        "Từ các phân tích trên, có thể nhận thấy rằng bốn thách thức T1–T4 không độc lập hoàn toàn: quá tải bình luận "
        "(T1) và thiếu ghi nhận tín hiệu phi ngôn ngữ (T3) cùng làm tăng gánh nặng lên người phát; thiếu thử trực quan "
        "(T2) và triển khai rời rạc (T4) cùng làm giảm hiệu quả của quy trình hỗ trợ bán hàng. Trong phạm vi luận văn "
        "này, bốn nhóm thách thức được ánh xạ trực tiếp sang nhóm yêu cầu hệ thống tại Bảng 3.1 và được giải quyết "
        "bằng hai trụ cột công nghệ — thị giác máy tính và xử lý ngôn ngữ tự nhiên — trình bày lần lượt tại mục 2.3 "
        "và 2.4.",
    )
    add_body(
        doc,
        "Kết luận, ý nghĩa của mục 2.2 đối với hệ thống được đề xuất là cung cấp khung thách thức có thể kiểm chứng: "
        "mỗi chức năng trong nguyên mẫu phải trả lời được nó giảm quá tải nào, bổ sung thử trực quan ra sao, ghi nhận "
        "tín hiệu phi ngôn ngữ thế nào và nằm ở đâu trong quy trình phiên thống nhất. Mục 2.3 mở đầu trụ cột thị giác "
        "máy tính, hướng tới T2 và T3 từ góc lý thuyết xử lý hình ảnh và thực tế tăng cường.",
    )


def add_section_2_3(doc):
    add_heading(doc, "2.3. Thị giác máy tính trong hỗ trợ bán hàng")
    add_body(
        doc,
        "Sau khi xác định thách thức nghiệp vụ, luận văn chuyển sang trụ cột thị giác máy tính (Computer Vision — CV). "
        "Trụ cột này không nhằm xây dựng một hệ thống thị giác máy tính tổng quát, mà phục vụ ba nhu cầu cụ thể trong "
        "phiên bán hàng: ghi nhận người xem, nhận dạng cử chỉ quan tâm và hỗ trợ thử sản phẩm trực quan trên trình "
        "duyệt. Cách tiếp cận này giữ đúng trục luận văn — bài toán trước, công nghệ sau.",
    )

    add_heading(doc, "2.3.1. Tổng quan thị giác máy tính")
    add_body(
        doc,
        "Theo Szeliski (2022), thị giác máy tính bao gồm chuỗi xử lý từ thu nhận hình ảnh, trích xuất đặc trưng "
        "đến suy luận cấp cao (nhận dạng đối tượng, ước lượng hình dạng, theo dõi). Trong thương mại trực tuyến, "
        "luồng video từ camera người dùng tạo nguồn dữ liệu liên tục; các thuật toán CV cần đáp ứng yêu cầu thời "
        "gian thực tương đối để không làm gián đoạn trải nghiệm xem phiên phát trực tiếp. Cần phân biệt xử lý "
        "hình ảnh thuần túy (phát hiện, nhận dạng) với trải nghiệm thực tế tăng cường (Augmented Reality — AR), "
        "vốn chồng lớp nội dung ảo lên khung hình thực. Trong phạm vi luận văn, cả hai hướng đều phục vụ quy trình "
        "Hình 2.1, không tách rời khỏi ngữ cảnh phiên.",
    )

    add_heading(doc, "2.3.2. Nhận diện khuôn mặt")
    add_body(
        doc,
        "Nhận diện khuôn mặt trong bối cảnh phiên phát trực tiếp phục vụ mục đích ghi nhận người xem đã đăng ký, "
        "theo dõi sự xuất hiện hoặc rời khung hình, hỗ trợ người phát nắm bắt nhóm khán giả quen thuộc. "
        "Các phương pháp hiện đại thường dựa trên trích xuất vector đặc trưng (embedding) từ mạng nơ-ron sâu, "
        "sau đó so khớp với cơ sở dữ liệu đã đăng ký. Deng và cộng sự (2019) đề xuất hàm mất ArcFace, cải thiện "
        "khả năng phân tách embedding giữa các danh tính, là nền tảng cho nhiều bộ công cụ nhận diện khuôn mặt "
        "thực tế. Quy trình khái niệm gồm: thu nhận khung hình, phát hiện khuôn mặt, trích xuất embedding, so sánh "
        "ngưỡng với mẫu đã lưu. Chi tiết triển khai trên nguyên mẫu sẽ trình bày tại mục 3.3.",
    )

    add_heading(doc, "2.3.3. Nhận dạng cử chỉ")
    add_body(
        doc,
        "Cử chỉ tay là tín hiệu phi ngôn ngữ phổ biến trong tương tác trực tuyến, trực tiếp liên quan thách thức T3. "
        "Nhận dạng cử chỉ trong thời gian thực yêu cầu ước lượng vị trí khớp tay (hand landmarks) từ luồng video. "
        "Zhang và cộng sự (2020) trình bày mô hình MediaPipe Hands, tối ưu cho thiết bị đa dạng và phù hợp triển "
        "khai trên trình duyệt thông qua WebAssembly. Khung MediaPipe (Lugaresi và cộng sự, 2019) cung cấp pipeline "
        "xử lý luồng video thống nhất, được áp dụng cho cả face mesh và hand tracking. Trong hỗ trợ bán hàng, cử chỉ "
        "có thể được diễn giải thành sự kiện nghiệp vụ (ví dụ thể hiện quan tâm) để bổ sung cho bình luận văn bản, "
        "thay vì yêu cầu người phát tự quan sát toàn bộ khung hình.",
    )

    add_heading(doc, "2.3.4. Thực tế tăng cường")
    add_body(
        doc,
        "Thực tế tăng cường chồng mô hình hoặc hình ảnh sản phẩm lên video camera người dùng, cho phép thử trực quan "
        "mà không cần chạm vào sản phẩm vật lý. Azuma và cộng sự (2001) định nghĩa AR là sự kết hợp giữa thế giới thực "
        "và nội dung ảo tương tác theo thời gian thực. Hoffmann và Mai (2022) tổng hợp bằng chứng cho thấy AR shopping "
        "giảm độ không chắc chắn khi mua hàng trực tuyến — phù hợp với thách thức T2. Khi tích hợp vào phiên phát "
        "trực tiếp, AR cần liên kết với sản phẩm đang được giới thiệu (sản phẩm ghim) và không tách rời khỏi luồng "
        "tương tác bình luận. Thiết kế quy trình xử lý tín hiệu hình ảnh sẽ được trình bày tại mục 3.3 và Hình 3.2.",
    )
    add_body(
        doc,
        "Dựa trên các nghiên cứu trên, có thể nhận thấy rằng thị giác máy tính trong luận văn đóng vai trò cầu nối "
        "giữa thách thức T2, T3 và quy trình Hình 2.1: nhận diện khuôn mặt và cử chỉ ghi nhận tín hiệu phi ngôn ngữ; "
        "thực tế tăng cường bổ sung trải nghiệm thử sản phẩm trực quan. Trong phạm vi luận văn này, các kỹ thuật CV "
        "được chọn theo tiêu chí triển khai được trên trình duyệt web và phù hợp thời gian thực tương đối, không nhằm "
        "so sánh benchmark học thuật độc lập.",
    )
    add_body(
        doc,
        "Kết luận, ý nghĩa của mục 2.3 đối với hệ thống được đề xuất là xác lập lý thuyết cho nhánh hình ảnh trong "
        "quy trình hỗ trợ bán hàng: từ thu nhận tín hiệu camera đến sự kiện nghiệp vụ phục vụ người phát. Thách thức "
        "T1 (quá tải bình luận) chủ yếu thuộc chiều ngôn ngữ; mục 2.4 trình bày trụ cột xử lý ngôn ngữ tự nhiên — "
        "phân loại ý định bình luận và xác định sản phẩm được nhắc đến — bổ sung cho nhánh CV đã nêu, hướng tới hệ "
        "thống tích hợp hai trụ cột theo thiết kế Chương 3.",
    )


def add_section_2_4(doc):
    add_heading(doc, "2.4. Xử lý ngôn ngữ tự nhiên trong bình luận người xem")
    add_body(
        doc,
        "Thách thức T1 (quá tải bình luận) chủ yếu liên quan chiều ngôn ngữ: người phát phải đọc hiểu và phản hồi "
        "hàng loạt bình luận ngắn, không chuẩn hóa, xuất hiện theo thời gian thực. Mục 2.4 trình bày trụ cột xử lý "
        "ngôn ngữ tự nhiên (Natural Language Processing — NLP) phục vụ hai bài toán tách biệt trong cùng phiên: "
        "(1) phân loại ý định bình luận; (2) xác định sản phẩm được nhắc đến. Việc tách bạch hai bài toán tránh nhầm "
        "lẫn giữa “người xem muốn gì” và “người xem đang nói về sản phẩm nào” — nền tảng cho thiết kế mục 3.4.",
    )

    add_heading(doc, "2.4.1. Tổng quan xử lý ngôn ngữ tự nhiên")
    add_body(
        doc,
        "Theo Jurafsky và Martin (2023), pipeline NLP điển hình gồm tiền xử lý văn bản, biểu diễn ngữ nghĩa và mô hình "
        "suy luận (phân loại, gán nhãn). Với bình luận thương mại, tiền xử lý cần chuẩn hóa biến thể tiếng Việt không "
        "dấu, viết tắt và từ lóng phổ biến trong môi trường trực tuyến. Kết quả phân loại ý định là đầu vào cho trợ lý "
        "bán hàng quyết định phản hồi tự động, chuyển tiếp hay bỏ qua — sẽ trình bày tại mục 2.5.",
    )

    add_heading(doc, "2.4.2. Phân loại ý định bình luận")
    add_body(
        doc,
        "Phân loại ý định (intent classification) gán mỗi bình luận vào một lớp phản ánh mục đích giao tiếp của người "
        "xem. Các mô hình dựa trên BERT (Devlin và cộng sự, 2019) và biến thể đa ngôn ngữ đã trở thành chuẩn cho bài "
        "toán phân loại câu ngắn trong hội thoại và dịch vụ khách hàng (Chen và cộng sự, 2019). Trong phạm vi luận văn, "
        "tập ý định được xác định theo nhu cầu nghiệp vụ livestream commerce, gồm mười một lớp như Bảng 2.1.",
    )
    table_21_rows = [
        ("1", "ASK_PRICE", "Hỏi giá sản phẩm", "“giá bao nhiêu?”, “bn tiền shop?”"),
        ("2", "ASK_VARIANT", "Hỏi biến thể (size, màu…)", "“còn size M không?”, “còn màu đen không?”"),
        ("3", "ASK_STOCK", "Hỏi tồn kho", "“còn hàng không shop?”"),
        ("4", "ASK_LINK", "Xin liên kết mua", "“cho link mua với”, “xin link”"),
        ("5", "ASK_SHIPPING", "Hỏi vận chuyển", "“ship Hà Nội mấy ngày?”"),
        ("6", "ASK_PROMOTION", "Hỏi khuyến mãi", "“có giảm giá không?”"),
        ("7", "PRODUCT_INFO", "Hỏi thông tin sản phẩm", "“son này hợp da dầu không?”"),
        ("8", "PURCHASE_INTENT", "Ý định chốt đơn", "“chốt 2 cái”, “mình mua”"),
        ("9", "CHITCHAT", "Trò chuyện xã giao", "“live vui quá”, “hello shop”"),
        ("10", "COMPLAINT", "Khiếu nại, phàn nàn", "“hàng lỗi quá shop ơi”"),
        ("11", "SPAM_TOXIC", "Spam hoặc nội dung độc hại", "Quảng cáo lạ, ngôn từ xúc phạm"),
    ]
    add_table(doc, ["STT", "Mã ý định", "Mô tả", "Ví dụ bình luận"], table_21_rows)
    add_caption(doc, "Bảng 2.1. Phân loại mười một ý định bình luận trong phiên phát trực tiếp")
    add_body(doc, "Nguồn: Tổng hợp tác giả theo nhu cầu nghiệp vụ hỗ trợ bán hàng.")

    add_heading(doc, "2.4.3. Mô hình ngôn ngữ PhoBERT")
    add_body(
        doc,
        "PhoBERT do Nguyen và Nguyen (2020) công bố là mô hình ngôn ngữ tiền huấn luyện cho tiếng Việt, dựa trên "
        "kiến trúc BERT và huấn luyện trên corpus tiếng Việt quy mô lớn. Mô hình phù hợp bài toán phân loại câu ngắn "
        "bằng cách fine-tune lớp phân loại trên đầu ra token [CLS]. Trong luận văn, PhoBERT đóng vai trò công cụ "
        "nhận dạng ý định bình luận tiếng Việt; chi tiết huấn luyện, siêu tham số và tập dữ liệu được trình bày tại "
        "Phụ lục và đánh giá tại mục 4.5. Mục này chỉ nêu khái niệm: đầu vào là chuỗi bình luận đã tiền xử lý, đầu ra "
        "là nhãn ý định thuộc Bảng 2.1 cùng độ tin cậy.",
    )

    add_heading(doc, "2.4.4. Phương pháp TF-IDF trong nhận diện sản phẩm được nhắc đến")
    add_body(
        doc,
        "Khác với phân loại ý định, xác định sản phẩm được nhắc trong bình luận là bài toán khớp văn bản với mô tả "
        "trong danh mục. TF-IDF (Term Frequency — Inverse Document Frequency) là phương pháp vector hóa văn bản cổ "
        "điển, trong đó trọng số phản ánh tần suất thuật ngữ trong tài liệu và độ hiếm trong tập tài liệu (Salton và "
        "Buckley, 1988). Cosine similarity giữa vector bình luận và vector mô tả sản phẩm cho phép xếp hạng sản phẩm "
        "liên quan mà không cần mô hình ngôn ngữ sâu. TF-IDF bổ sung cho PhoBERT ở chiều sản phẩm: PhoBERT trả lời "
        "“ý định là gì”, TF-IDF hỗ trợ “đang nói về sản phẩm nào”.",
    )
    table_23_rows = [
        ("ASK_PRICE", "ASK_PRICE", "AUTO_REPLY_SUGGESTED"),
        ("ASK_VARIANT", "ASK_SIZE", "AUTO_REPLY_SUGGESTED"),
        ("ASK_STOCK", "ASK_STOCK", "AUTO_REPLY_SUGGESTED"),
        ("ASK_LINK", "ASK_LINK", "AUTO_REPLY_SUGGESTED"),
        ("ASK_SHIPPING", "ASK_SHIPPING", "AUTO_REPLY_SUGGESTED"),
        ("ASK_PROMOTION", "ASK_PROMOTION", "AUTO_REPLY_SUGGESTED"),
        ("PRODUCT_INFO", "ASK_PRODUCT_INFO", "AUTO_REPLY_SUGGESTED"),
        ("PURCHASE_INTENT", "PURCHASE_INTENT", "ESCALATE_TO_HOST"),
        ("CHITCHAT", "UNKNOWN", "IGNORE"),
        ("COMPLAINT", "ASK_PRODUCT_INFO", "ESCALATE_TO_HOST"),
        ("SPAM_TOXIC", "UNKNOWN", "IGNORE"),
    ]
    add_table(doc, ["Nhãn mô hình (ML)", "Nhãn hệ thống", "Hành động gợi ý"], table_23_rows)
    add_caption(doc, "Bảng 2.3. Ánh xạ nhãn ý định mô hình sang nhãn và hành động hệ thống")
    add_body(doc, "Nguồn: Tổng hợp tác giả theo thiết kế trợ lý bán hàng (chi tiết tại mục 3.4).")

    add_body(
        doc,
        "Từ các nghiên cứu và thiết kế trên, có thể nhận thấy rằng xử lý ngôn ngữ trong luận văn không gộp thành một "
        "pipeline duy nhất: nhánh ý định (PhoBERT) và nhánh sản phẩm (TF-IDF) phục vụ câu hỏi nghiệp vụ khác nhau, "
        "rồi hội tụ tại trợ lý bán hàng. Trong phạm vi luận văn này, taxonomy mười một lớp là đóng góp ứng dụng ở tầng "
        "nghiệp vụ — bổ sung cho công cụ NLP có sẵn, phù hợp bối cảnh bình luận tiếng Việt trong phiên bán hàng.",
    )
    add_body(
        doc,
        "Kết luận, ý nghĩa của mục 2.4 đối với hệ thống được đề xuất là cung cấp lý thuyết và phân loại cho chiều xử "
        "lý bình luận — giảm trực tiếp gánh nặng T1 bằng cách tự động nhận diện ý định và sản phẩm liên quan. Mục 2.5 "
        "trình bày khái niệm trợ lý bán hàng thông minh: thành phần chuyển kết quả NLP (và tín hiệu từ mục 2.3) thành "
        "phản hồi hoặc đề xuất hành động cho người phát theo quy trình Hình 2.1.",
    )


def add_section_2_5(doc):
    add_heading(doc, "2.5. Trợ lý bán hàng thông minh")
    add_body(
        doc,
        "Hai trụ cột thị giác máy tính và xử lý ngôn ngữ tự nhiên tạo ra tín hiệu và phân loại, nhưng chưa tự thành "
        "quy trình hỗ trợ bán hàng nếu thiếu lớp điều phối nghiệp vụ. Trợ lý bán hàng thông minh (Intelligent Sales "
        "Assistant) là thành phần trung gian chuyển kết quả phân tích bình luận, tín hiệu hình ảnh và ngữ cảnh sản phẩm "
        "thành phản hồi hoặc đề xuất hành động cho người phát trực tiếp, như vị trí thể hiện ở Hình 2.1.",
    )
    add_body(
        doc,
        "Følstad và Skjuve (2019) chỉ ra rằng người dùng đánh giá chatbot qua khả năng giải quyết vấn đề và tốc độ "
        "phản hồi đúng ngữ cảnh; khi chatbot thất bại, họ kỳ vọng chuyển tiếp nhanh sang nhân viên thật. Luận điểm đó "
        "ủng hộ thiết kế trợ lý thiên thực dụng trong phiên livestream: không phải mọi bình luận đều nên được trả lời "
        "tự động, đặc biệt với chốt đơn và khiếu nại.",
    )
    add_body(
        doc,
        "Về mặt lý thuyết, trợ lý cần phân biệt ít nhất bốn kiểu xử lý: (1) đề xuất phản hồi tự động khi ý định và sản "
        "phẩm đủ rõ ràng; (2) chuyển vào hàng đợi cần người phát xử lý; (3) leo thang ưu tiên cao khi có ý định mua "
        "hoặc khiếu nại; (4) bỏ qua bình luận không liên quan hoặc spam. Bảng 2.3 thể hiện ánh xạ từ nhãn mô hình sang "
        "hành động gợi ý (AUTO_REPLY_SUGGESTED, ESCALATE_TO_HOST, IGNORE) — khái niệm hóa trước thiết kế luật nghiệp "
        "vụ tại mục 3.5.",
    )
    add_body(
        doc,
        "Trợ lý cũng cần gắn phản hồi với sản phẩm đã xác định (từ tên nhắc trực tiếp, khớp ngữ nghĩa hoặc sản phẩm "
        "ghim), tránh trả lời chung chung không phù hợp ngữ cảnh phiên. Quy trình khái niệm tại Hình 2.1 thể hiện vị "
        "trí trợ lý sau bước phân tích ý định và sản phẩm; luồng end-to-end sẽ được mô tả chi tiết tại Hình 3.5.",
    )
    add_body(
        doc,
        "Dựa trên các phân tích trên, có thể nhận thấy rằng trợ lý bán hàng là điểm hội tụ của hai trụ cột công nghệ "
        "và mô hình vận hành mục 2.1: thay vì thay thế hoàn toàn người phát, trợ lý giảm tải các tác vụ lặp lại có thể "
        "tự động hóa và leo thang các tình huống nhạy cảm. Trong phạm vi luận văn này, trợ lý được thiết kế theo nguyên "
        "tắc “hỗ trợ quyết định”, phù hợp conversational commerce trong bối cảnh bình luận công khai đồng thời.",
    )
    add_body(
        doc,
        "Kết luận, ý nghĩa của mục 2.5 đối với hệ thống được đề xuất là hoàn thiện khung lý thuyết giữa phân tích "
        "(§2.3–2.4) và thiết kế (Ch.3): mọi tín hiệu đi qua trợ lý trước khi trở thành phản hồi hoặc thao tác mua hàng. "
        "Mục 2.6 sẽ đặt nguyên mẫu đề xuất trong bối cảnh nghiên cứu liên quan và chỉ ra khoảng trống tích hợp (T4) "
        "mà luận văn hướng tới.",
    )


def add_section_2_6(doc):
    add_heading(doc, "2.6. Các nghiên cứu liên quan")
    add_body(
        doc,
        "Các mục 2.1–2.5 đã trình bày cơ sở lý thuyết theo trục nghiệp vụ: mô hình vận hành phiên bán hàng, "
        "thách thức T1–T4, hai trụ cột thị giác máy tính và xử lý ngôn ngữ tự nhiên, cùng khái niệm trợ lý "
        "bán hàng. Mục 2.6 đặt nguyên mẫu đề xuất trong bối cảnh nghiên cứu liên quan, so sánh có hệ thống "
        "với các công trình tiêu biểu và chỉ ra khoảng trống mà luận văn hướng tới lấp đầy.",
    )
    add_body(
        doc,
        "Phương pháp so sánh chọn năm công trình đại diện cho năm hướng bổ sung cho nhau: thương mại phiên "
        "phát trực tiếp (Wongkitrungrueng và Assarut, 2020); trải nghiệm chatbot và leo thang (Følstad và Skjuve, "
        "2019); mô hình ngôn ngữ tiếng Việt (Nguyen và Nguyen, 2020 — PhoBERT); thực tế tăng cường trong mua "
        "sắm (Hoffmann và Mai, 2022); trí tuệ nhân tạo đa phương thức cho thương mại (Yu và cộng sự, 2022 — "
        "CommerceMM). Tiêu chí chọn không phải liệt kê đầy đủ toàn bộ literature mà là bao phủ các lớp công "
        "nghệ và bối cảnh nghiệp vụ liên quan trực tiếp tới thiết kế Chương 3.",
    )
    table_22_rows = [
        (
            "Wongkitrungrueng & Assarut, 2020",
            "Mô hình tin cậy và engagement trong livestream commerce; củng cố vai trò tương tác thời gian thực",
            "Góc nhìn người mua; không lớp hỗ trợ kỹ thuật cho người phát",
            "§2.1 · T1 · G2 → Đ2 (quy trình hỗ trợ người phát)",
        ),
        (
            "Følstad & Skjuve, 2019",
            "UX chatbot: hiệu quả trả lời và leo thang quan trọng hơn “giống người”",
            "Bối cảnh chat tĩnh; không livestream / bình luận đồng thời",
            "§2.5 · AUTO/ESCALATE · G2 → Đ2",
        ),
        (
            "Nguyen & Nguyen, 2020 (PhoBERT)",
            "Mô hình ngôn ngữ tiếng Việt mạnh trên task NLP chuẩn",
            "Không domain bình luận bán hàng livestream",
            "§2.4.3 · G3 → Đ3 (taxonomy 11 lớp)",
        ),
        (
            "Hoffmann & Mai, 2022",
            "Review AR shopping: virtual try-on giảm rủi ro cảm nhận",
            "AR chủ yếu app / kênh tách rời phiên phát",
            "§2.3.4 · T2 · G1, G2 → Đ1, Đ2 (AR trong phiên)",
        ),
        (
            "Yu et al., 2022 (CommerceMM)",
            "AI đa phương thức catalog quy mô nền tảng",
            "Hướng marketplace tĩnh; không mô hình phiên cho người phát",
            "§2.6 đối chứng · G1 → Đ1 (POC tích hợp phiên)",
        ),
    ]
    add_table(doc, ["Công trình", "Đóng góp", "Hạn chế", "Liên hệ với đề tài"], table_22_rows)
    add_caption(doc, "Bảng 2.2. So sánh các nghiên cứu liên quan tiêu biểu")
    add_body(doc, "Nguồn: Tổng hợp tác giả theo khảo sát literature (Literature Freeze v1.0).")

    add_body(
        doc,
        "Từ Bảng 2.2, có thể nhóm ba khoảng trống nghiên cứu chính. G1 — khoảng trống tích hợp: AR, trợ lý, "
        "NLP và AI catalog thường được nghiên cứu tách biệt; ít mô tả nguyên mẫu gộp các lớp trong một phiên "
        "phát trực tiếp (liên quan T4). G2 — khoảng trống bối cảnh phiên livestream: công nghệ hướng tới chat "
        "tĩnh, ứng dụng AR hoặc marketplace, trong khi literature livestream thiên về người mua, ít mô tả hỗ trợ "
        "người phát trong luồng video và bình luận đồng thời (T1, T2). G3 — khoảng trống domain NLP tiếng Việt: "
        "PhoBERT mạnh trên task chuẩn nhưng chưa có taxonomy bình luận bán hàng livestream tiếng Việt (T1 ở chiều "
        "ngôn ngữ).",
    )
    add_body(
        doc,
        "Luận văn đề xuất ba đóng góp tương ứng, bổ sung cho nhau: Đ1 — nguyên mẫu tích hợp trên trình duyệt web, "
        "đồng thời có lớp thị giác máy tính, lớp NLP và trợ lý bán hàng (trả lời G1, T4); Đ2 — quy trình hỗ trợ "
        "trong phiên: bình luận → phân tích → ý định → trợ lý (AUTO/ESCALATE) → thử sản phẩm ghim, gắn phiên phát "
        "trực tiếp (trả lời G2, T1, T2); Đ3 — taxonomy mười một lớp ý định bình luận bán hàng livestream tiếng Việt, "
        "fine-tune PhoBERT trên domain và kết hợp nhánh khớp sản phẩm TF-IDF (trả lời G3). Đ1 nhấn kiến trúc tích "
        "hợp; Đ2 nhấn quy trình nghiệp vụ phiên — hai đóng góp không trùng lặp.",
    )
    add_body(
        doc,
        "Hình 2.3 trực quan hóa khoảng trống: bên trái là các giải pháp rời rạc (chatbot, AR, nhận diện, nền tảng "
        "phát sóng); giữa là nhãn khoảng trống; bên phải là hệ thống hỗ trợ bán hàng thông minh tích hợp hai trụ cột "
        "trong một nguyên mẫu phiên. Hình bổ sung cho Bảng 2.2 ở mức khái niệm, không thay thế bằng chứng chi tiết "
        "từng công trình.",
    )
    add_figure(
        doc,
        "hinh_2_3_khoang_trong_nghien_cuu.png",
        "Hình 2.3. Khoảng trống nghiên cứu: giải pháp rời rạc và hệ thống tích hợp đề xuất",
    )

    add_body(
        doc,
        "Dựa trên các phân tích trên, có thể nhận thấy rằng điểm khác biệt của luận văn không nằm ở việc đề xuất "
        "một thuật toán CV hay NLP mới, mà ở việc mô tả và hiện thực nguyên mẫu hội tụ phục vụ quy trình Hình 2.1 "
        "trong bối cảnh phiên bán hàng tiếng Việt. Trong phạm vi luận văn này, so sánh literature khẳng định tính "
        "cần thiết của thiết kế Chương 3 trước khi trình bày hiện thực và đánh giá tại Chương 4.",
    )
    add_body(
        doc,
        "Kết luận, ý nghĩa của mục 2.6 đối với hệ thống được đề xuất là neo đóng góp nghiên cứu vào bằng chứng "
        "literature có căn cứ: mỗi chức năng trong nguyên mẫu trả lời một phần G1–G3. Mục 2.7 tóm tắt công nghệ triển "
        "khai được chọn để hiện thực hai trụ cột đã trình bày, trước khi chuyển sang phân tích yêu cầu và kiến trúc "
        "tại Chương 3.",
    )


def add_section_2_7(doc):
    add_heading(doc, "2.7. Công nghệ sử dụng trong hệ thống")
    add_body(
        doc,
        "Sau khi xác định khoảng trống nghiên cứu, mục 2.7 tóm tắt công nghệ được chọn để hiện thực nguyên mẫu — "
        "không nhằm liệt kê stack triển khai độc lập với bài toán, mà ánh xạ từng nhóm công nghệ vào vai trò trong "
        "hai trụ cột thị giác máy tính và xử lý ngôn ngữ tự nhiên đã trình bày ở mục 2.3 và 2.4. Chi tiết kênh giao "
        "tiếp và phiên bản phần mềm được trình bày tại Bảng 3.4 (mục 3.7).",
    )
    add_body(
        doc,
        "Trụ cột thị giác máy tính: (1) thử sản phẩm trực quan bằng thực tế tăng cường phía trình duyệt, chồng "
        "hình ảnh sản phẩm lên luồng camera người xem; (2) nhận diện khuôn mặt dựa trên InsightFace (embedding "
        "ArcFace) để ghi nhận người xem đã đăng ký; (3) nhận dạng cử chỉ qua MediaPipe Hands, biến tín hiệu phi "
        "ngôn ngữ thành sự kiện nghiệp vụ. Ba nhóm này phục vụ trực tiếp thách thức T2 và T3.",
    )
    add_body(
        doc,
        "Trụ cột xử lý ngôn ngữ tự nhiên: (1) PhoBERT (Nguyen và Nguyen, 2020) fine-tune cho phân loại mười một "
        "lớp ý định bình luận (§2.4.3); (2) TF-IDF kết hợp cosine similarity để khớp bình luận với mô tả sản phẩm "
        "trong danh mục (§2.4.4); (3) luật nghiệp vụ dự phòng khi dịch vụ phân loại không sẵn sàng, duy trì vận "
        "hành phiên. Ba nhóm phục vụ thách thức T1 và bổ sung cho trợ lý bán hàng (§2.5).",
    )
    add_body(
        doc,
        "Hình 2.2 thể hiện hai trụ cột đối xứng: mỗi trụ gồm ba khối vai trò nghiệp vụ — không phải sơ đồ luồng "
        "thời gian. Ghi chú triển khai (lớp giao diện React, lớp máy chủ FastAPI, container Docker) chỉ xuất hiện "
        "ở mức phụ, vì chúng phục vụ triển khai nguyên mẫu thử nghiệm chứ không phải đóng góp lý thuyết của luận văn.",
    )
    add_figure(
        doc,
        "hinh_2_2_hai_tru_cot_cong_nghe.png",
        "Hình 2.2. Hai trụ cột công nghệ của hệ thống hỗ trợ bán hàng thông minh",
    )

    add_body(
        doc,
        "Về triển khai, nguyên mẫu chọn kiến trúc hoạt động trên trình duyệt web: lớp giao diện (React/TypeScript) "
        "xử lý tương tác người dùng, thực tế tăng cường và một phần nhận dạng cử chỉ; lớp máy chủ (FastAPI/Python) "
        "điều phối trò chuyện thời gian thực, đăng ký khuôn mặt, chuyển tiếp bình luận tới dịch vụ PhoBERT và ghi "
        "nhận sự kiện tương tác. Docker hỗ trợ triển khai thử nghiệm cục bộ. Lựa chọn này bám yêu cầu tương tác thời "
        "gian thực trong phiên (§2.1) và phạm vi nguyên mẫu nghiên cứu đã nêu tại Chương 1 — không bao gồm cơ sở "
        "dữ liệu đơn hàng bền vững hay cổng thanh toán thật.",
    )
    add_body(
        doc,
        "Dựa trên các phân tích trên, có thể nhận thấy rằng công nghệ trong luận văn được tổ chức theo bài toán hỗ "
        "trợ bán hàng: PhoBERT và TF-IDF là công cụ NLP, MediaPipe và InsightFace là công cụ CV, React/FastAPI/Docker "
        "là phương tiện hiện thực — không đảo ngược thứ tự ưu tiên nghiệp vụ → lý thuyết → công nghệ đã thiết lập "
        "từ đầu chương.",
    )
    add_body(
        doc,
        "Kết luận, ý nghĩa của mục 2.7 đối với hệ thống được đề xuất là cung cấp bức tranh công nghệ ngắn gọn, "
        "sẵn sàng chuyển sang thiết kế: Chương 3 sẽ mô tả yêu cầu (§3.1), kiến trúc tổng thể (§3.2) và thiết kế chi "
        "tiết từng luồng xử lý — không lặp lại giảng giải lý thuyết đã có ở §2.3–2.4.",
    )


def add_chapter2_conclusion(doc):
    add_heading(doc, "Kết luận chương")
    add_body(
        doc,
        "Chương 2 đã xây dựng cơ sở lý thuyết và nghiên cứu liên quan theo trục nghiệp vụ phiên bán hàng trực tiếp. "
        "Mục 2.1 mô tả mô hình tương tác và quy trình hỗ trợ (Hình 2.1); mục 2.2 formal hóa bốn thách thức T1–T4; "
        "mục 2.3 và 2.4 trình bày hai trụ cột thị giác máy tính và xử lý ngôn ngữ tự nhiên; mục 2.5 định nghĩa trợ "
        "lý bán hàng như điểm hội tụ; mục 2.6 so sánh năm công trình tiêu biểu (Bảng 2.2, Hình 2.3) và chỉ ra khoảng "
        "trống G1–G3 cùng đóng góp Đ1–Đ3; mục 2.7 tóm tắt công nghệ triển khai (Hình 2.2).",
    )
    add_body(
        doc,
        "Từ khung lý thuyết này, Chương 3 chuyển sang phân tích yêu cầu hệ thống và thiết kế kiến trúc nguyên mẫu: "
        "ánh xạ T1–T4 sang chức năng cụ thể (Bảng 3.1), mô tả kiến trúc tổng thể (Hình 3.1) và thiết kế chi tiết "
        "từng luồng xử lý hình ảnh, bình luận, trợ lý bán hàng và quy trình mua hàng mô phỏng. Chương 4 sẽ trình bày "
        "hiện thực và đánh giá các thiết kế đó bằng kịch bản thử nghiệm và bằng chứng đo lường.",
    )


def build(sync_manuscript: bool = False):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    add_chapter_intro(doc)
    add_section_2_1(doc)
    add_section_2_2(doc)
    add_section_2_3(doc)
    add_section_2_4(doc)
    add_section_2_5(doc)
    add_section_2_6(doc)
    add_section_2_7(doc)
    add_chapter2_conclusion(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")

    if sync_manuscript:
        OUT_REVIEW.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(OUT_REVIEW))
        print(f"Synced manuscript: {OUT_REVIEW}")
    else:
        print("Manuscript package: not synced (use --sync-manuscript when section locked)")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Chapter 2 DOCX")
    parser.add_argument(
        "--sync-manuscript",
        action="store_true",
        help="Copy to thesis_review_package/manuscript/ (only when section is locked)",
    )
    args = parser.parse_args()
    build(sync_manuscript=args.sync_manuscript)
