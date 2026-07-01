# -*- coding: utf-8 -*-
"""Generate Chapter 3 DOCX — Execution Mode (batch writing).

Batch 1: §3.1 + §3.2
Batch 2: §3.3 + §3.4 + §3.5 (later)
Batch 3: §3.6 + §3.7 + §3.8 (later)
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = ROOT / "docs" / "thesis-figures"
OUT = ROOT / "Document" / "Tieu luan - Smart Livestream - Chuong 3.docx"
OUT_REVIEW = ROOT / "thesis_review_package" / "manuscript" / OUT.name

COMPLETED_BATCHES = ["1", "2", "3"]
CHAPTER_LOCKED = True
BATCH_STATUS = "Chương 3 đã khóa (Batch 1–3)"
REVIEW_NOTE = "docs/chapter3/REVIEW_NOTE_BATCH3_3_6_3_8.md"


def set_run_font(run, name="Times New Roman", size=13, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_para(doc, text, size=13, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, size=13, bold=True):
    return add_para(doc, text, size=size, bold=bold, space_after=8)


def add_body(doc, text):
    return add_para(doc, text, size=13, bold=False, space_after=6)


def add_caption(doc, text):
    p = add_para(doc, text, size=13, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    for run in p.runs:
        run.italic = True
    return p


def add_figure(doc, filename, caption, width_cm=14):
    path = FIG_DIR / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=12)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=12)
    doc.add_paragraph()
    return table


def add_chapter_intro(doc):
    add_heading(doc, "CHƯƠNG 3", size=14)
    add_heading(doc, "PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", size=14)
    add_body(
        doc,
        "Chương này chuyển khung lý thuyết Chương 2 sang thiết kế nguyên mẫu hệ thống hỗ trợ bán hàng trong "
        "phiên phát trực tiếp. Nội dung được tổ chức theo quy trình nghiệp vụ: phân tích yêu cầu, kiến trúc "
        "tổng thể, thiết kế xử lý hình ảnh và bình luận, trợ lý bán hàng, quy trình mua hàng mô phỏng và kiến "
        "trúc triển khai. Chương mô tả thiết kế; hiện thực và đánh giá được trình bày tại Chương 4.",
    )


def add_section_3_1(doc):
    add_heading(doc, "3.1. Phân tích yêu cầu hệ thống")
    add_body(
        doc,
        "Thiết kế hệ thống xuất phát từ thách thức T1–T4 (mục 2.2) và quy trình hỗ trợ bán hàng (Hình 2.1). "
        "Mục 3.1 xác định tác nhân tham gia, nhóm yêu cầu chức năng và phi chức năng, ánh xạ trực tiếp sang "
        "chức năng nguyên mẫu — làm cơ sở cho kiến trúc tổng thể tại mục 3.2.",
    )

    add_heading(doc, "3.1.1. Tác nhân và bối cảnh sử dụng")
    add_body(
        doc,
        "Nguyên mẫu có hai tác nhân chính trong cùng một phiên phát trực tiếp mô phỏng. Người xem theo dõi "
        "luồng video, gửi bình luận, sử dụng thử sản phẩm trên camera, thực hiện cử chỉ quan tâm và thao tác "
        "mua hàng mô phỏng. Người phát trực tiếp trình bày sản phẩm, ghim sản phẩm đang giới thiệu, theo dõi "
        "trợ lý bán hàng và bảng sự kiện AI, quyết định chốt đơn hoặc can thiệp khi bình luận được leo thang. "
        "Hệ thống hỗ trợ đóng vai trò trung gian: ghi nhận đa kênh tín hiệu, phân tích và đề xuất phản hồi — "
        "không thay thế hoàn toàn quyết định của người phát.",
    )

    add_heading(doc, "3.1.2. Yêu cầu chức năng và phi chức năng")
    add_body(
        doc,
        "Bảng 3.1 tổng hợp yêu cầu chức năng (YC-CN) và phi chức năng (YC-PNC), kèm thách thức T1–T4 tương ứng. "
        "Các yêu cầu được rút ra từ kịch bản vận hành phiên bán hàng và phạm vi nguyên mẫu đã nêu tại Chương 1.",
    )
    table_31_rows = [
        ("YC-CN-01", "Chức năng", "Phân loại ý định bình luận tiếng Việt", "T1", "Giảm quá tải đọc bình luận"),
        ("YC-CN-02", "Chức năng", "Xác định sản phẩm được nhắc trong bình luận", "T1", "Gắn phản hồi đúng ngữ cảnh"),
        ("YC-CN-03", "Chức năng", "Thử sản phẩm trực quan (thực tế tăng cường) trên trình duyệt", "T2", "Hỗ trợ quyết định mua"),
        ("YC-CN-04", "Chức năng", "Đăng ký và nhận diện khuôn mặt người xem", "T3", "Ghi nhận khán giả quen"),
        ("YC-CN-05", "Chức năng", "Nhận dạng cử chỉ và hiển thị trên bảng sự kiện AI", "T3", "Ghi nhận tín hiệu phi ngôn ngữ"),
        ("YC-CN-06", "Chức năng", "Trợ lý bán hàng: đề xuất phản hồi, leo thang, bỏ qua", "T1, T4", "Hội tụ CV + NLP trong phiên"),
        ("YC-CN-07", "Chức năng", "Trò chuyện trực tuyến đồng bộ trong phiên", "T1", "Luồng tương tác công khai đồng thời"),
        ("YC-CN-08", "Chức năng", "Mô phỏng giỏ hàng và quy trình thanh toán", "T4", "Khép kín quy trình Hình 2.1"),
        ("YC-PNC-01", "Phi CN", "Vận hành trên trình duyệt web, không yêu cầu cài đặt ứng dụng gốc", "—", "Phù hợp phiên trực tuyến"),
        ("YC-PNC-02", "Phi CN", "Phản hồi tương đối thời gian thực cho bình luận và sự kiện AI", "T1", "Giữ trải nghiệm phiên"),
        ("YC-PNC-03", "Phi CN", "Dữ liệu phiên lưu tạm; mất khi khởi động lại dịch vụ", "—", "Phạm vi nguyên mẫu nghiên cứu"),
        ("YC-PNC-04", "Phi CN", "Duy trì vận hành khi dịch vụ PhoBERT không sẵn sàng (luật dự phòng)", "T1", "Giảm phụ thuộc một điểm lỗi"),
    ]
    add_table(doc, ["Mã", "Loại", "Yêu cầu", "Thách thức", "Mục đích nghiệp vụ"], table_31_rows)
    add_caption(doc, "Bảng 3.1. Yêu cầu chức năng và phi chức năng của hệ thống")
    add_body(doc, "Nguồn: Tổng hợp tác giả theo thách thức T1–T4 và phạm vi nguyên mẫu (Chương 1).")

    add_body(
        doc,
        "Dựa trên Bảng 3.1, có thể nhận thấy rằng yêu cầu chức năng bao phủ đủ hai trụ cột công nghệ và lớp "
        "trợ lý: YC-CN-01, 02, 06, 07 thuộc chiều ngôn ngữ và điều phối; YC-CN-03, 04, 05 thuộc chiều hình ảnh; "
        "YC-CN-08 khép quy trình mua hàng. Yêu cầu phi chức năng giới hạn phạm vi nghiên cứu — không yêu cầu hạ "
        "tầng sản xuất quy mô lớn hay lưu trữ đơn hàng bền vững.",
    )
    add_body(
        doc,
        "Kết luận, ý nghĩa của mục 3.1 đối với hệ thống được đề xuất là cung cấp danh sách yêu cầu có thể "
        "kiểm chứng: mỗi chức năng hiện thực tại Chương 4 phải map về ít nhất một mã YC-CN. Mục 3.2 mô tả kiến "
        "trúc tổng thể thỏa mãn các yêu cầu trên trong một phiên thống nhất.",
    )


def add_section_3_2(doc):
    add_heading(doc, "3.2. Kiến trúc tổng thể")
    add_body(
        doc,
        "Kiến trúc tổng thể mô tả cách các thành phần hỗ trợ bán hàng phối hợp trong một phiên — theo góc "
        "nhìn nghiệp vụ, không liệt kê theo tên module mã nguồn. Hình 3.1 thể hiện luồng từ người xem xuống "
        "người phát trực tiếp, với hệ thống hỗ trợ bán hàng thông minh ở trung tâm và thành phần AI phục vụ "
        "phân tích hai trụ cột.",
    )
    add_body(
        doc,
        "Luồng nghiệp vụ xương sống: người xem tương tác qua bình luận, camera và thao tác mua hàng → lớp giao "
        "diện phiên (xem video, trò chuyện, thử sản phẩm, giỏ hàng) → hệ thống hỗ trợ ghi nhận tín hiệu hình "
        "ảnh (bảng sự kiện AI), phân tích bình luận (ý định và sản phẩm), điều phối trợ lý bán hàng → thành phần "
        "AI thực hiện nhận dạng cử chỉ/khuôn mặt, thực tế tăng cường, PhoBERT và TF-IDF → kết quả hỗ trợ trả về "
        "cho người phát (gợi ý phản hồi, leo thang, sự kiện) và cho người xem (phản hồi trò chuyện, thử sản phẩm). "
        "Luồng này khớp quy trình Hình 2.1 và đóng góp Đ2 (quy trình phiên).",
    )
    add_body(
        doc,
        "Về phân tách trách nhiệm: lớp giao diện đảm nhiệm tương tác người dùng và một phần xử lý hình ảnh phía "
        "trình duyệt (thực tế tăng cường, cử chỉ); lớp máy chủ điều phối trò chuyện thời gian thực, lưu tạm hồ "
        "sơ khuôn mặt, chuyển tiếp bình luận tới dịch vụ phân loại ý định và ghi nhận sự kiện tương tác. Thiết "
        "kế chi tiết từng luồng được trình bày lần lượt tại mục 3.3 (hình ảnh), 3.4 (bình luận), 3.5 (trợ lý) "
        "và 3.6 (mua hàng); kiến trúc triển khai vật lý (cổng, giao thức) tại mục 3.7 và Bảng 3.4.",
    )
    add_figure(doc, "hinh_3_1_kien_truc_tong_the.png", "Hình 3.1. Kiến trúc tổng thể hệ thống hỗ trợ bán hàng (góc nhìn nghiệp vụ)")

    add_body(
        doc,
        "Dựa trên kiến trúc trên, có thể nhận thấy rằng nguyên mẫu hiện thực đóng góp Đ1: một giao diện phiên "
        "tích hợp đồng thời hai trụ cột — không tách AR, chatbot và nhận diện thành công cụ độc lập. Trong phạm "
        "vi luận văn này, kiến trúc tổng thể là khung tham chiếu cho mọi thiết kế chi tiết và đánh giá end-to-end "
        "tại Chương 4.",
    )
    add_body(
        doc,
        "Kết luận, ý nghĩa của mục 3.2 đối với hệ thống được đề xuất là xác định ranh giới và luồng dữ liệu "
        "nghiệp vụ trước khi đi vào thuật toán cụ thể. Mục 3.3 mở đầu thiết kế chi tiết nhánh thị giác máy tính "
        "(Hình 3.2), tiếp nối trụ cột đã trình bày tại mục 2.3.",
    )


def add_section_3_3(doc):
    add_heading(doc, "3.3. Thiết kế xử lý tín hiệu hình ảnh")
    add_body(
        doc,
        "Mục 3.3 chi tiết hóa trụ cột thị giác máy tính trong kiến trúc Hình 3.1. Thiết kế tách ba nhánh "
        "song song từ cùng nguồn camera người xem: thử sản phẩm trực quan, nhận diện khuôn mặt và nhận "
        "dạng cử chỉ — mỗi nhánh phục vụ yêu cầu YC-CN-03, 04, 05.",
    )
    add_body(
        doc,
        "Nhánh thử sản phẩm: luồng video camera được xử lý phía trình duyệt bằng MediaPipe (điểm mốc "
        "khuôn mặt) và canvas thực tế tăng cường. Hiệu ứng thử (kính, son, bộ lọc…) được ánh xạ theo "
        "sản phẩm đang ghim trong phiên — khi người phát đổi sản phẩm ghim, hiệu ứng thay đổi tương ứng. "
        "Nhánh này giải quyết trực tiếp T2 mà không tách AR khỏi phiên phát.",
    )
    add_body(
        doc,
        "Nhánh nhận diện khuôn mặt: khung hình được gửi tới lớp máy chủ, trích xuất embedding bằng "
        "InsightFace và so khớp với mẫu đã đăng ký trong phiên. Kết quả hỗ trợ ghi nhận người xem quen "
        "thuộc xuất hiện trên khung hình. Nhánh nhận dạng cử chỉ: luồng video được phân tích (MediaPipe "
        "Hands trên lớp máy chủ hoặc trình duyệt tùy luồng), cử chỉ như giơ tay hoặc giơ ngón cái được "
        "chuyển thành sự kiện nghiệp vụ với cơ chế cooldown chống lặp spam.",
    )
    add_body(
        doc,
        "Ba nhánh hội tụ tại bảng sự kiện AI trên giao diện phiên: người phát theo dõi tín hiệu phi ngôn "
        "ngữ và nhận diện mà không phải quan sát toàn bộ khung hình thủ công. Hình 3.2 mô tả luồng chi tiết.",
    )
    add_figure(doc, "hinh_3_2_quy_trinh_xu_ly_tin_hieu_hinh_anh.png", "Hình 3.2. Quy trình xử lý tín hiệu hình ảnh trong phiên phát trực tiếp")
    add_body(
        doc,
        "Kết luận, thiết kế xử lý hình ảnh đảm bảo tín hiệu CV phục vụ quy trình hỗ trợ bán hàng (Hình 2.1), "
        "không tồn tại như demo nhận diện độc lập. Mục 3.4 chuyển sang trụ cột xử lý bình luận.",
    )


def add_section_3_4(doc):
    add_heading(doc, "3.4. Thiết kế xử lý bình luận")
    add_body(
        doc,
        "Xử lý bình luận được thiết kế thành hai nhánh độc lập — phân loại ý định và xác định sản phẩm — "
        "chỉ hội tụ tại trợ lý bán hàng (mục 3.5). Thiết kế này tránh nhầm lẫn “người xem muốn gì” với "
        "“người xem đang nói về sản phẩm nào”, đã nêu tại mục 2.4.",
    )

    add_heading(doc, "3.4.1. Nhánh phân loại ý định")
    add_body(
        doc,
        "Luồng xử lý: bình luận người xem → chuẩn hóa văn bản (bỏ dấu, chuẩn hóa viết tắt) → phân loại "
        "sơ bộ bằng luật/biểu thức → gọi mô hình PhoBERT qua lớp máy chủ chuyển tiếp → kết hợp lai "
        "(ưu tiên mô hình khi độ tin cậy đạt ngưỡng) → ánh xạ sang mười một lớp ý định hệ thống (Bảng 2.1). "
        "Khi dịch vụ PhoBERT không phản hồi, hệ thống chuyển sang luật dự phòng (YC-PNC-04). Hình 3.3 "
        "mô tả luồng này — không bao gồm TF-IDF hay danh mục sản phẩm.",
    )
    add_figure(doc, "hinh_3_3_quy_trinh_xac_dinh_y_dinh.png", "Hình 3.3. Quy trình xác định ý định bình luận")

    add_heading(doc, "3.4.2. Nhánh xác định sản phẩm")
    add_body(
        doc,
        "Luồng xử lý: bình luận đã chuẩn hóa → tìm kiếm TF-IDF trên danh mục mười sản phẩm mô phỏng → "
        "áp dụng thứ tự ưu tiên (Bảng 3.2) → sản phẩm được chọn kèm độ tin cậy và nguồn xác định. Nhánh "
        "này không sử dụng PhoBERT. Khi nhiều sản phẩm cùng loại có điểm gần nhau, hệ thống đánh dấu mơ hồ "
        "và có thể sinh câu hỏi làm rõ.",
    )
    table_32_rows = [
        ("1", "Tên sản phẩm / mã được nhắc trực tiếp", "mentioned_product", "Điểm khớp tên hoặc thẻ cao nhất"),
        ("2", "Tìm kiếm ngữ nghĩa TF-IDF", "semantic_search", "Cosine similarity trên mô tả catalog"),
        ("3", "Khớp theo loại sản phẩm (từ khóa danh mục)", "category_match", "Ví dụ: “son”, “kính”"),
        ("4", "Sản phẩm đang ghim trong phiên", "pinned_product", "Mặc định khi câu không cụ thể"),
        ("5", "Làm rõ / mơ hồ", "ambiguous", "Nhiều ứng viên cùng điểm — hỏi thêm người xem"),
    ]
    add_table(doc, ["Bước", "Tiêu chí", "Nguồn xác định", "Ghi chú"], table_32_rows)
    add_caption(doc, "Bảng 3.2. Thứ tự ưu tiên xác định sản phẩm được nhắc trong bình luận")
    add_body(doc, "Nguồn: Tổng hợp tác giả theo thiết kế bộ giải quyết sản phẩm trong nguyên mẫu.")

    add_figure(doc, "hinh_3_4_quy_trinh_xac_dinh_san_pham.png", "Hình 3.4. Quy trình xác định sản phẩm được nhắc trong bình luận")

    add_body(
        doc,
        "Kết luận, thiết kế xử lý bình luận tách bạch hai bài toán NLP phục vụ YC-CN-01 và YC-CN-02. "
        "Mục 3.5 mô tả cách trợ lý bán hàng hợp nhất kết quả hai nhánh.",
    )


def add_section_3_5(doc):
    add_heading(doc, "3.5. Thiết kế trợ lý bán hàng thông minh")
    add_body(
        doc,
        "Trợ lý bán hàng là điểm hội tập duy nhất của nhánh ý định (Hình 3.3) và nhánh sản phẩm (Hình 3.4). "
        "Đầu vào gồm: ý định cuối cùng và độ tin cậy; sản phẩm được chọn và nguồn xác định; ngữ cảnh phiên "
        "(sản phẩm ghim, cấu hình tự động phản hồi). Lớp luật nghiệp vụ quyết định hành động theo Bảng 2.3: "
        "AUTO_REPLY_SUGGESTED (đề xuất phản hồi tự động), ESCALATE_TO_HOST (leo thang cho người phát), "
        "IGNORE (bỏ qua bình luận không liên quan hoặc spam).",
    )
    add_body(
        doc,
        "Sau khi chọn hành động, bộ sinh phản hồi mẫu tạo câu trả lời gắn với sản phẩm đã xác định — tránh "
        "phản hồi chung chung. Phản hồi hiển thị trên kênh trò chuyện phiên (công khai) và/hoặc bảng trợ lý "
        "bán hàng (hỗ trợ người phát). Với ý định PURCHASE_INTENT hoặc COMPLAINT, trợ lý ưu tiên leo thang "
        "thay vì trả lời tự động, phù hợp khuyến nghị Følstad và Skjuve (2019) về kỳ vọng chuyển tiếp sang "
        "người thật.",
    )
    add_body(
        doc,
        "Trợ lý cũng gợi ý thao tác mua hàng (thêm giỏ, mở quy trình thanh toán) khi ý định và sản phẩm đủ "
        "rõ ràng — liên kết sang mục 3.6. Hình 3.5 mô tả luồng end-to-end của trợ lý.",
    )
    add_figure(doc, "hinh_3_5_quy_trinh_ho_tro_ban_hang.png", "Hình 3.5. Quy trình hỗ trợ bán hàng thông minh (trợ lý bán hàng)")
    add_body(
        doc,
        "Kết luận, trợ lý hiện thực YC-CN-06 và quy trình Hình 2.1 ở tầng thiết kế: giảm tải T1, hội tụ hai "
        "trụ cột trong một phiên (T4). Mục 3.6 mô tả quy trình mua hàng mô phỏng khi người xem thực hiện thao tác mua.",
    )


def add_section_3_6(doc):
    add_heading(doc, "3.6. Thiết kế quy trình mua hàng mô phỏng")
    add_body(
        doc,
        "Quy trình mua hàng mô phỏng khép vòng trải nghiệm phiên: người xem chọn sản phẩm (từ danh mục, "
        "sản phẩm ghim hoặc gợi ý trợ lý) → thêm vào giỏ hàng (số lượng, tạm tính) → mở biểu mẫu thanh toán "
        "mô phỏng → chọn một trong hai nhánh: thanh toán khi nhận hàng (COD) hoặc thanh toán điện tử mô phỏng "
        "(ví dụ quét mã QR giả lập) → nhận mã đơn hàng (tiền tố LS-) và trạng thái đơn (cod_confirmed hoặc paid).",
    )
    add_body(
        doc,
        "Dữ liệu giỏ hàng và đơn hàng chỉ tồn tại trong phiên trình duyệt (localStorage hoặc bộ nhớ tạm), "
        "không ghi vào cơ sở dữ liệu bền vững — phù hợp phạm vi nghiên cứu Chương 1. Thông tin trên biểu mẫu "
        "thanh toán là dữ liệu mô phỏng, không kết nối cổng thanh toán thật. Thiết kế này hiện thực YC-CN-08.",
    )
    add_figure(doc, "hinh_3_6_quy_trinh_mua_hang.png", "Hình 3.6. Quy trình mua hàng mô phỏng trong nguyên mẫu")
    add_body(
        doc,
        "Kết luận, quy trình mua hàng mô phỏng cho phép đánh giá end-to-end luồng bán hàng trong phiên mà "
        "không vượt phạm vi đề tài. Mục 3.7 mô tả kiến trúc triển khai vật lý của các thành phần.",
    )


def add_section_3_7(doc):
    add_heading(doc, "3.7. Kiến trúc triển khai")
    add_body(
        doc,
        "Kiến trúc triển khai bổ sung cho Hình 3.1 ở tầng vật lý: nguyên mẫu chạy trên môi trường cục bộ "
        "hoặc Docker Compose, gồm container lớp giao diện (Vite/React), container lớp máy chủ (FastAPI) và "
        "dịch vụ PhoBERT (Python) phục vụ phân loại ý định. Trình duyệt người dùng truy cập lớp giao diện; "
        "lớp giao diện gọi REST API và WebSocket tới lớp máy chủ theo Bảng 3.4.",
    )
    table_34_rows = [
        ("Lớp giao diện (React)", "HTTP", "Tải ứng dụng phiên · thực tế tăng cường phía trình duyệt"),
        ("Trò chuyện phiên", "WebSocket", "/ws/chat/{room_id} — đồng bộ bình luận đa thiết bị"),
        ("Luồng thời gian thực AI", "WebSocket", "/ws/realtime — sự kiện bổ sung"),
        ("API máy chủ", "REST", "/api/health · /api/frame · /api/interaction-events · đăng ký khuôn mặt"),
        ("Phân loại ý định", "REST", "/api/nlp/predict-intent — dịch vụ PhoBERT"),
        ("Luồng video MJPEG", "HTTP", "/api/video-feed — xử lý khung hình phía máy chủ"),
    ]
    add_table(doc, ["Thành phần", "Giao thức / kênh", "Mục đích"], table_34_rows)
    add_caption(doc, "Bảng 3.4. Kênh giao tiếp và dịch vụ trong kiến trúc triển khai")
    add_body(doc, "Nguồn: Tổng hợp tác giả theo cấu hình nguyên mẫu (mục 4.1).")

    add_body(
        doc,
        "Hình 3.7 mô tả sơ đồ triển khai: trình duyệt ↔ lớp giao diện ↔ lớp máy chủ ↔ dịch vụ PhoBERT. "
        "Chi tiết phiên bản phần mềm và phần cứng thử nghiệm được trình bày tại Bảng 4.1 (Chương 4).",
    )
    add_figure(doc, "hinh_3_7_kien_truc_trien_khai.png", "Hình 3.7. Kiến trúc triển khai nguyên mẫu")
    add_body(
        doc,
        "Kết luận, kiến trúc triển khai đảm bảo các thiết kế nghiệp vụ §3.3–3.6 có thể vận hành thử nghiệm "
        "tích hợp trong một phiên. Mục 3.8 tóm tắt chương.",
    )


def add_section_3_8(doc):
    add_heading(doc, "3.8. Tóm tắt chương")
    add_body(
        doc,
        "Chương 3 đã chuyển khung lý thuyết Chương 2 sang thiết kế nguyên mẫu hệ thống hỗ trợ bán hàng "
        "trong phiên phát trực tiếp. Mục 3.1 xác định yêu cầu và ánh xạ T1–T4 (Bảng 3.1); mục 3.2 mô tả "
        "kiến trúc tổng thể (Hình 3.1); mục 3.3 thiết kế xử lý tín hiệu hình ảnh (Hình 3.2); mục 3.4 thiết "
        "kế xử lý bình luận hai nhánh (Hình 3.3, 3.4, Bảng 3.2); mục 3.5 thiết kế trợ lý bán hàng (Hình 3.5); "
        "mục 3.6 thiết kế quy trình mua hàng mô phỏng (Hình 3.6); mục 3.7 trình bày kiến trúc triển khai "
        "(Hình 3.7, Bảng 3.4).",
    )
    add_body(
        doc,
        "Thiết kế thống nhất quy trình Hình 2.1 và đóng góp nghiên cứu Đ1–Đ3: một giao diện phiên tích hợp "
        "hai trụ cột, luồng hỗ trợ người phát và taxonomy bình luận tiếng Việt. Chương 4 sẽ trình bày hiện "
        "thực các thiết kế trên và đánh giá bằng kịch bản thử nghiệm, ảnh chụp giao diện và chỉ số đo lường.",
    )


def build(sync_manuscript: bool = False, batch: int = 3):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    add_chapter_intro(doc)
    if batch >= 1:
        add_section_3_1(doc)
        add_section_3_2(doc)
    if batch >= 2:
        add_section_3_3(doc)
        add_section_3_4(doc)
        add_section_3_5(doc)
    if batch >= 3:
        add_section_3_6(doc)
        add_section_3_7(doc)
        add_section_3_8(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")

    if sync_manuscript:
        OUT_REVIEW.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(OUT_REVIEW))
        print(f"Synced manuscript: {OUT_REVIEW}")
    else:
        print("Manuscript package: not synced (use --sync-manuscript when batch locked)")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Chapter 3 DOCX")
    parser.add_argument("--batch", type=int, default=1, help="Highest batch to include (1–3)")
    parser.add_argument(
        "--sync-manuscript",
        action="store_true",
        help="Copy to thesis_review_package/manuscript/ when batch locked",
    )
    args = parser.parse_args()
    build(sync_manuscript=args.sync_manuscript, batch=args.batch)
